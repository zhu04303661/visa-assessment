#!/usr/bin/env python3
"""批量读取docx文件，提取文本内容并输出为结构化格式。

用法:
  # 读取单个文件
  python scripts/read_docx.py path/to/file.docx

  # 读取整个目录下所有docx
  python scripts/read_docx.py path/to/directory/

  # 输出为JSON格式（含段落级结构）
  python scripts/read_docx.py --json path/to/file.docx

  # 输出到文件
  python scripts/read_docx.py path/to/directory/ -o output.json
"""
import argparse
import json
import os
import sys
from pathlib import Path

from docx import Document


def read_single_docx(filepath: str) -> dict:
    """读取单个docx文件，返回结构化数据。"""
    doc = Document(filepath)
    paragraphs = []
    tables_data = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        paragraphs.append({
            "text": text,
            "style": para.style.name if para.style else "Normal",
            "bold": any(run.bold for run in para.runs if run.bold),
        })

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            tables_data.append(rows)

    return {
        "file": os.path.basename(filepath),
        "path": str(filepath),
        "paragraph_count": len(paragraphs),
        "table_count": len(tables_data),
        "paragraphs": paragraphs,
        "tables": tables_data,
        "full_text": "\n".join(p["text"] for p in paragraphs),
    }


def read_directory(dirpath: str, recursive: bool = False) -> list:
    """读取目录下所有docx文件。"""
    results = []
    p = Path(dirpath)
    pattern = "**/*.docx" if recursive else "*.docx"
    for f in sorted(p.glob(pattern)):
        if f.name.startswith("~$"):
            continue
        try:
            results.append(read_single_docx(str(f)))
        except Exception as e:
            results.append({"file": f.name, "path": str(f), "error": str(e)})
    return results


def format_text_output(data) -> str:
    """将结构化数据格式化为可读文本。"""
    if isinstance(data, list):
        parts = []
        for item in data:
            parts.append(format_single_text(item))
        return "\n\n".join(parts)
    return format_single_text(data)


def format_single_text(item: dict) -> str:
    """格式化单个文件的输出。"""
    if "error" in item:
        return f"=== {item['file']} ===\n[ERROR] {item['error']}"

    lines = [
        f"{'=' * 60}",
        f"文件: {item['file']}",
        f"段落数: {item['paragraph_count']}  |  表格数: {item['table_count']}",
        f"{'=' * 60}",
        "",
        item["full_text"],
    ]

    if item["tables"]:
        lines.append(f"\n--- 表格数据 ({item['table_count']}个) ---")
        for i, table in enumerate(item["tables"], 1):
            lines.append(f"\n[表格{i}]")
            for row in table:
                lines.append(" | ".join(row))

    return "\n".join(lines)


def main():
    parser = argparse.ArgumentParser(description="读取docx文件提取结构化文本")
    parser.add_argument("path", help="docx文件路径或包含docx的目录路径")
    parser.add_argument("--json", action="store_true", help="输出JSON格式")
    parser.add_argument("-o", "--output", help="输出到文件")
    parser.add_argument("-r", "--recursive", action="store_true", help="递归搜索子目录")
    args = parser.parse_args()

    target = Path(args.path)
    if target.is_file():
        data = read_single_docx(str(target))
    elif target.is_dir():
        data = read_directory(str(target), recursive=args.recursive)
    else:
        print(f"路径不存在: {args.path}", file=sys.stderr)
        sys.exit(1)

    if args.json:
        output = json.dumps(data, ensure_ascii=False, indent=2)
    else:
        output = format_text_output(data)

    if args.output:
        Path(args.output).write_text(output, encoding="utf-8")
        print(f"已输出至: {args.output}")
    else:
        print(output)


if __name__ == "__main__":
    main()
