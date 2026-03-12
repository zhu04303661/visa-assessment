#!/usr/bin/env python3
"""生成推荐人CV的docx文件。

CV格式结构：
- 标题 "个人简历"，居中加粗
- 基本信息区（姓名、性别、邮箱、单位地址等，逐行"键：值"排列，不使用表格）
- 职业经历（每段经历独立列出：职位加粗、单位、地点、时间，按时间倒序）
- 教育背景（学位加粗、学校、时间）
- 专业领域（逐行列出）
- 荣誉（逐行列出）
- 主要成就（逐条列出详细描述）
- 信息出处（尾部附出处链接）

整体风格：扁平简洁，不使用表格，段落紧凑。

JSON输入格式见底部的 EXAMPLE_INPUT。
"""
import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

FONT_NAME = "仿宋"


def _set_run_font(run, size=11):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)


def setup_document(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.5


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    _set_run_font(run, size=16)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return p


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    _set_run_font(run, size=12)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return p


def add_line(doc, text, size=11, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.74)
    run = p.add_run(text)
    run.bold = bold
    _set_run_font(run, size=size)
    return p


def add_key_value_line(doc, key, value, size=11):
    p = doc.add_paragraph()
    run_key = p.add_run(f"{key}：")
    run_key.bold = True
    _set_run_font(run_key, size=size)
    run_val = p.add_run(value)
    _set_run_font(run_val, size=size)
    return p


def add_experience_entry(doc, title, org, location="", period=""):
    p_title = doc.add_paragraph()
    run = p_title.add_run(title)
    run.bold = True
    _set_run_font(run, size=11)

    if org:
        add_key_value_line(doc, "单位", org)
    if location:
        add_key_value_line(doc, "地点", location)
    if period:
        add_key_value_line(doc, "时间", period)


def generate_cv(data: dict, output_path: str):
    doc = Document()
    setup_document(doc)

    add_title(doc, "个人简历")

    # --- 基本信息区 ---
    basic = data.get("basic_info", {})
    name = data.get("name", basic.get("name", ""))

    info_pairs = []
    if name:
        gender = basic.get("gender", "")
        if gender:
            info_pairs.append(("姓名", f"{name}　　　　性别：{gender}"))
        else:
            info_pairs.append(("姓名", name))
    if basic.get("email"):
        info_pairs.append(("电子邮箱", basic["email"]))
    if basic.get("address"):
        info_pairs.append(("单位地址", basic["address"]))
    if basic.get("website"):
        info_pairs.append(("个人主页", basic["website"]))

    for key, val in info_pairs:
        add_key_value_line(doc, key, val)

    # --- 职业经历 ---
    if data.get("work_experience"):
        add_section_heading(doc, "职业经历")
        for exp in data["work_experience"]:
            add_experience_entry(
                doc,
                title=exp.get("title", ""),
                org=exp.get("organization", ""),
                location=exp.get("location", ""),
                period=exp.get("period", ""),
            )
            if exp.get("description"):
                add_line(doc, exp["description"], size=10, indent=True)
            doc.add_paragraph()

    # --- 教育背景 ---
    if data.get("education"):
        add_section_heading(doc, "教育背景")
        for edu in data["education"]:
            degree = edu.get("degree", "")
            institution = edu.get("institution", "")
            year = edu.get("year", edu.get("period", ""))
            detail = edu.get("detail", "")

            p = doc.add_paragraph()
            run = p.add_run(degree)
            run.bold = True
            _set_run_font(run, size=11)

            if institution:
                add_key_value_line(doc, "学校", institution)
            if year:
                add_key_value_line(doc, "时间", year)
            if detail:
                add_line(doc, detail, size=10, indent=True)

    # --- 专业领域 ---
    if data.get("research_areas"):
        add_section_heading(doc, "专业领域")
        for area in data["research_areas"]:
            add_line(doc, area)

    # --- 荣誉 ---
    if data.get("honors"):
        add_section_heading(doc, "荣誉")
        for honor in data["honors"]:
            add_line(doc, honor)

    # --- 主要成就 ---
    if data.get("achievements"):
        add_section_heading(doc, "主要成就")
        for ach in data["achievements"]:
            add_line(doc, ach)
            doc.add_paragraph()

    # --- 信息出处 ---
    if data.get("sources"):
        add_section_heading(doc, "信息出处")
        for i, src in enumerate(data["sources"], 1):
            label = src.get("label", f"来源{i}")
            url = src.get("url", "")
            add_line(doc, f"[{i}] {label}: {url}", size=9)

    output_dir = Path(output_path).parent
    output_dir.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description="从JSON生成推荐人CV的docx文件")
    parser.add_argument("input", help="CV数据JSON文件路径")
    parser.add_argument("-o", "--output", required=True, help="输出docx文件路径")
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"文件不存在: {args.input}", file=sys.stderr)
        sys.exit(1)

    data = json.loads(input_path.read_text(encoding="utf-8"))
    output_dir = Path(args.output).parent
    output_dir.mkdir(parents=True, exist_ok=True)

    result = generate_cv(data, args.output)
    print(f"CV已生成: {result}")


EXAMPLE_INPUT = """{
  "name": "推荐人姓名",
  "basic_info": {
    "name": "推荐人姓名",
    "gender": "男",
    "email": "email@example.com",
    "address": "单位地址",
    "website": ""
  },
  "work_experience": [
    {
      "title": "职务名称",
      "organization": "单位名称",
      "location": "城市，国家",
      "period": "起始年份至今"
    }
  ],
  "education": [
    {
      "degree": "学位",
      "institution": "院校名称",
      "year": "毕业年份"
    }
  ],
  "research_areas": [
    "专业领域1",
    "专业领域2"
  ],
  "honors": [
    "荣誉1",
    "荣誉2"
  ],
  "achievements": [
    "主要成就描述1。",
    "主要成就描述2。"
  ],
  "sources": [
    {"label": "来源名称", "url": "https://example.com/..."}
  ]
}"""


if __name__ == "__main__":
    main()
