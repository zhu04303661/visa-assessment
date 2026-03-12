#!/usr/bin/env python3
"""生成GTV推荐信的docx文件。

格式规范：
- 仅一行中文标题"推荐信"
- 无独立日期/收件人/主题头部块
- 正文为连续自然段落，无编号章节
- 段落内 **加粗文本** 自动转为Word加粗格式
- 仅在信尾有一个中文签名落款块（姓名、职务、单位、邮箱、电话、地址、日期）
- 不包含"出处列表"（出处保留在构思思路文件中）

JSON输入格式见底部的 EXAMPLE_INPUT。
"""
import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

FONT_NAME = "仿宋"


def _set_run_font(run, size=12):
    """统一设置 run 的中西文字体为仿宋。"""
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)


def setup_document(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(4)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    _set_run_font(run, size=16)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return p


def add_rich_paragraph(doc, text, indent_first_line=True, font_size=12):
    """添加段落，自动将 **text** 转为Word加粗格式。"""
    p = doc.add_paragraph()
    if indent_first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)

    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            _set_run_font(run, size=font_size)
        else:
            run = p.add_run(part)
            _set_run_font(run, size=font_size)
    return p


def add_signature_block(doc, data):
    doc.add_paragraph()

    sig_lines = []
    if data.get("name"):
        sig_lines.append(data["name"])
    if data.get("titles"):
        sig_lines.extend(data["titles"])
    elif data.get("title"):
        sig_lines.append(data["title"])
    if data.get("organization"):
        sig_lines.append(data["organization"])
    if data.get("email"):
        sig_lines.append(f"电子邮箱：{data['email']}")
    if data.get("phone"):
        sig_lines.append(f"联系电话：{data['phone']}")
    if data.get("address"):
        sig_lines.append(data["address"])
    if data.get("date"):
        sig_lines.append(f"日期：{data['date']}")

    for line in sig_lines:
        if not line:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(line)
        _set_run_font(run, size=11)


SKIP_PATTERNS = [
    r'^[\*\*]*推荐人信息出处',
    r'^[\*\*]*申请人证据材料出处',
    r'^\d+\.\s*(百度百科|维基百科)',
    r'^-\s*(MC\d|OC\d|核心事实)',
    r'^https?://',
    r'^此致$',
    r'^敬礼$',
    r'^---$',
    r'^#\s',
    r'^\*\*日期\*\*',
    r'^\*\*致\*\*',
    r'^\*\*关于\*\*',
]


def should_skip(text: str) -> bool:
    for pat in SKIP_PATTERNS:
        if re.search(pat, text):
            return True
    return False


def is_section_header(text: str) -> bool:
    return bool(re.match(r'^(#{1,3}\s|[一二三四五六七八九十]+、)', text))


def clean_section_title(text: str) -> str:
    text = re.sub(r'^#{1,3}\s*', '', text)
    text = re.sub(r'^[一二三四五六七八九十]+、\s*', '', text)
    return text.strip()


def is_subsection_header(text: str) -> bool:
    return bool(re.match(r'^\d+\.\d+\s', text))


def clean_subsection_title(text: str) -> str:
    return re.sub(r'^\d+\.\d+\s*', '', text).strip()


def generate_letter(data: dict, output_path: str, template_path: str = None):
    if template_path and Path(template_path).exists():
        doc = Document(template_path)
        while doc.paragraphs:
            p = doc.paragraphs[0]
            p_element = p._element
            p_element.getparent().remove(p_element)
        for table in doc.tables:
            t_element = table._element
            t_element.getparent().remove(t_element)
    else:
        doc = Document()

    setup_document(doc)
    add_title(doc, "推荐信")

    doc.add_paragraph()

    paragraphs = data.get("paragraphs", [])

    if not paragraphs and data.get("sections"):
        for section in data["sections"]:
            sec_title = section.get("title", "")
            for para in section.get("paragraphs", []):
                paragraphs.append(para)

    seen_closing = False
    for text in paragraphs:
        text = text.strip()
        if not text:
            continue
        if should_skip(text):
            continue
        if text in ("此致", "敬礼"):
            if seen_closing:
                continue
            seen_closing = True
            continue
        if is_section_header(text):
            continue
        if is_subsection_header(text):
            continue
        if text.startswith("尊敬的评审委员会") and any("尊敬的" in p for p in paragraphs[:paragraphs.index(text)]):
            continue

        add_rich_paragraph(doc, text)

    add_signature_block(doc, data.get("signature", {}))

    output_dir = Path(output_path).parent
    output_dir.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description="从JSON生成推荐信docx文件（范本格式）")
    parser.add_argument("input", help="推荐信数据JSON文件路径")
    parser.add_argument("-o", "--output", required=True, help="输出docx文件路径")
    parser.add_argument("--template", help="可选：范本docx文件路径（仅复制页面设置）")
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"文件不存在: {args.input}", file=sys.stderr)
        sys.exit(1)

    data = json.loads(input_path.read_text(encoding="utf-8"))
    result = generate_letter(data, args.output, args.template)
    print(f"推荐信已生成: {result}")


EXAMPLE_INPUT = """{
  "paragraphs": [
    "本人[推荐人姓名]，[职务]...",
    "基于本人在[领域]的专业积累...",
    "我与[申请人]先生/女士结识于...",
    "[申请人]先生/女士最核心的贡献在于...",
    "我坚信[申请人]先生/女士完全符合英国全球人才签证标准。"
  ],
  "signature": {
    "name": "推荐人姓名",
    "titles": ["职务1", "职务2"],
    "organization": "单位名称",
    "email": "email@example.com",
    "address": "单位地址",
    "date": "2026年X月X日"
  }
}"""


if __name__ == "__main__":
    main()
