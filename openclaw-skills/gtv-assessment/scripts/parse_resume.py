#!/usr/bin/env python3
"""从 PDF / DOCX 简历文件中提取文本内容。

Usage:
    python3 parse_resume.py /path/to/resume.pdf
    python3 parse_resume.py /path/to/resume.docx

Output: 提取的纯文本内容输出到 stdout，供 AI 进一步解析结构化信息。
Metadata (file type, page count) 输出到 stderr。
"""
import sys
from pathlib import Path


def extract_from_pdf(file_path: str) -> tuple[str, dict]:
    """从 PDF 提取文本，返回 (text, metadata)。"""
    from pdfminer.high_level import extract_text
    from pdfminer.pdfpage import PDFPage

    text = extract_text(file_path) or ""
    metadata = {"file_type": "PDF", "page_count": None}

    try:
        with open(file_path, "rb") as f:
            pages = list(PDFPage.get_pages(f))
            metadata["page_count"] = len(pages)
    except Exception:
        pass

    return text.strip(), metadata


def extract_from_docx(file_path: str) -> tuple[str, dict]:
    """从 DOCX 提取文本（段落 + 表格），返回 (text, metadata)。"""
    from docx import Document

    doc = Document(file_path)
    parts = []

    # 段落
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)

    # 表格
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                parts.append(" | ".join(row_text))

    text = "\n".join(parts)
    metadata = {
        "file_type": "DOCX",
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
    }
    return text.strip(), metadata


def main() -> int:
    if len(sys.argv) < 2:
        print("Usage: python3 parse_resume.py <file_path>", file=sys.stderr)
        return 1

    file_path = sys.argv[1]
    path = Path(file_path)

    if not path.exists():
        print(f"Error: File not found: {file_path}", file=sys.stderr)
        return 1

    if not path.is_file():
        print(f"Error: Not a file: {file_path}", file=sys.stderr)
        return 1

    suffix = path.suffix.lower()

    if suffix == ".pdf":
        try:
            text, meta = extract_from_pdf(str(path))
        except ImportError as e:
            print(f"Error: pdfminer.six not installed: {e}", file=sys.stderr)
            return 1
        except Exception as e:
            print(f"Error: PDF parsing failed: {e}", file=sys.stderr)
            return 1

    elif suffix in (".docx",):
        try:
            text, meta = extract_from_docx(str(path))
        except ImportError as e:
            print(f"Error: python-docx not installed: {e}", file=sys.stderr)
            return 1
        except Exception as e:
            print(f"Error: DOCX parsing failed: {e}", file=sys.stderr)
            return 1

    elif suffix == ".doc":
        print(
            "Error: Legacy .doc format is not supported. Use .docx instead.",
            file=sys.stderr,
        )
        return 1

    else:
        print(
            f"Error: Unsupported format '{suffix}'. Use .pdf or .docx",
            file=sys.stderr,
        )
        return 1

    # Metadata to stderr
    meta_parts = [f"file_type={meta.get('file_type', 'unknown')}"]
    if meta.get("page_count") is not None:
        meta_parts.append(f"page_count={meta['page_count']}")
    if meta.get("paragraph_count") is not None:
        meta_parts.append(f"paragraph_count={meta['paragraph_count']}")
    if meta.get("table_count") is not None:
        meta_parts.append(f"table_count={meta['table_count']}")
    print(" ".join(meta_parts), file=sys.stderr)

    # Text to stdout
    print(text)
    return 0


if __name__ == "__main__":
    sys.exit(main())
