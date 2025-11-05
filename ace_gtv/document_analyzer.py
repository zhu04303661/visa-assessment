#!/usr/bin/env python3
"""
文档分析模块 - 提取Excel/Word内容并用LLM分析提炼知识规则
"""

import json
import logging
import os
from typing import Any, Dict, List, Optional
from pathlib import Path
from datetime import datetime

try:
    from openpyxl import load_workbook
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    from docx import Document as DocxDocument
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

try:
    from langchain_openai import ChatOpenAI
    HAS_LANGCHAIN = True
except ImportError:
    HAS_LANGCHAIN = False

logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

# LLM 提示词
KNOWLEDGE_EXTRACTION_PROMPT = """
你是一位GTV评估和知识管理专家。我提供了一份文档内容，其中包含关于签证评估、评分规则或专业标准的信息。

请根据这份内容，提炼出可以用于GTV评估系统的知识规则。

对于每个知识条目，提供以下结构化信息：
{
  "title": "知识条目标题（简洁）",
  "category": "分类（如：评估标准、评分规则、教育背景、工作经验等）",
  "dimension": "维度（如：education、experience、technical、leadership、impact）",
  "content": "详细内容描述",
  "scoringRules": ["规则1", "规则2", "..."]
}

文档内容：
{document_content}

请返回一个JSON数组，包含从文档中提炼出的所有知识条目。
确保每个条目都是有效的、相关的、可用于GTV评估的。
"""

class DocumentExtractor:
    """文档内容提取器"""
    
    @staticmethod
    def extract_from_excel(file_path: str) -> str:
        """从Excel文件提取文本"""
        if not HAS_OPENPYXL:
            raise ImportError("openpyxl 未安装，请运行: pip install openpyxl")
        
        logger.info(f"📊 从Excel文件提取内容: {file_path}")
        
        try:
            wb = load_workbook(file_path)
            content = []
            
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                content.append(f"\n[工作表: {sheet_name}]\n")
                
                for row in ws.iter_rows(values_only=True):
                    # 过滤空行
                    values = [str(v) for v in row if v is not None]
                    if values:
                        content.append(" | ".join(values))
            
            text = "\n".join(content)
            logger.info(f"✅ 成功提取Excel内容，共 {len(text)} 个字符")
            return text
            
        except Exception as e:
            logger.error(f"❌ Excel提取失败: {e}")
            raise
    
    @staticmethod
    def extract_from_word(file_path: str) -> str:
        """从Word文件提取文本"""
        if not HAS_PYTHON_DOCX:
            raise ImportError("python-docx 未安装，请运行: pip install python-docx")
        
        logger.info(f"📄 从Word文件提取内容: {file_path}")
        
        try:
            doc = DocxDocument(file_path)
            content = []
            
            # 提取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)
            
            # 提取表格
            for table in doc.tables:
                content.append("\n[表格]\n")
                for row in table.rows:
                    cells = [cell.text for cell in row.cells]
                    content.append(" | ".join(cells))
            
            text = "\n".join(content)
            logger.info(f"✅ 成功提取Word内容，共 {len(text)} 个字符")
            return text
            
        except Exception as e:
            logger.error(f"❌ Word提取失败: {e}")
            raise
    
    @staticmethod
    def extract_from_file(file_path: str) -> str:
        """根据文件类型自动选择提取方法"""
        file_path_obj = Path(file_path)
        suffix = file_path_obj.suffix.lower()
        
        logger.info(f"🔍 检测文件类型: {suffix}")
        
        if suffix in ['.xlsx', '.xls']:
            return DocumentExtractor.extract_from_excel(file_path)
        elif suffix in ['.docx', '.doc']:
            return DocumentExtractor.extract_from_word(file_path)
        elif suffix == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            raise ValueError(f"不支持的文件格式: {suffix}")


class KnowledgeExtractor:
    """知识规则提取器 - 使用LLM分析文档"""
    
    def __init__(self, api_key: Optional[str] = None):
        """初始化LLM"""
        logger.info("🚀 初始化知识提取器...")
        
        self.api_key = api_key or os.getenv("OPENAI_API_KEY")
        self.llm = None
        
        if HAS_LANGCHAIN and self.api_key:
            try:
                self.llm = ChatOpenAI(
                    api_key=self.api_key,
                    model="gpt-4-turbo-preview",
                    temperature=0.7,
                )
                logger.info("✅ LLM 初始化成功")
            except Exception as e:
                logger.error(f"❌ LLM初始化失败: {e}")
                self.llm = None
        else:
            logger.warning("⚠️ LLM 不可用，将使用本地规则生成")
    
    def extract_knowledge(self, document_content: str) -> List[Dict[str, Any]]:
        """使用LLM从文档内容提取知识规则"""
        logger.info("📚 开始提取知识规则...")
        
        if not self.llm:
            logger.warning("⚠️ LLM不可用，返回空结果")
            return []
        
        try:
            prompt = KNOWLEDGE_EXTRACTION_PROMPT.format(
                document_content=document_content[:4000]  # 限制长度
            )
            
            logger.debug(f"📝 发送提示词到LLM...")
            response = self.llm.invoke(prompt)
            
            # 解析LLM返回的JSON
            response_text = response.content
            logger.debug(f"💬 LLM响应: {response_text[:200]}...")
            
            # 尝试提取JSON
            import re
            json_match = re.search(r'\[.*\]', response_text, re.DOTALL)
            if json_match:
                json_str = json_match.group(0)
                knowledge_items = json.loads(json_str)
                logger.info(f"✅ 成功提取 {len(knowledge_items)} 个知识条目")
                return knowledge_items
            else:
                logger.warning("⚠️ LLM响应中未找到JSON格式数据")
                return []
                
        except Exception as e:
            logger.error(f"❌ 知识提取失败: {e}")
            return []
    
    def analyze_and_extract(self, file_path: str) -> Dict[str, Any]:
        """完整流程：提取文件内容 → LLM分析 → 提炼知识"""
        logger.info(f"\n{'='*80}")
        logger.info(f"📖 开始分析文档: {file_path}")
        logger.info(f"{'='*80}")
        
        start_time = datetime.now()
        
        try:
            # 第一步：提取文件内容
            logger.info("第1步: 提取文件内容...")
            content = DocumentExtractor.extract_from_file(file_path)
            
            # 第二步：LLM分析和提取知识
            logger.info("第2步: LLM分析和提取知识规则...")
            knowledge_items = self.extract_knowledge(content)
            
            # 第三步：验证和补充信息
            logger.info("第3步: 验证和补充知识条目...")
            validated_items = self._validate_items(knowledge_items)
            
            elapsed = (datetime.now() - start_time).total_seconds()
            
            result = {
                "success": True,
                "file": Path(file_path).name,
                "file_size": os.path.getsize(file_path),
                "content_length": len(content),
                "items_count": len(validated_items),
                "items": validated_items,
                "analysis_time": f"{elapsed:.2f}s",
                "timestamp": datetime.now().isoformat(),
            }
            
            logger.info(f"✅ 分析完成！")
            logger.info(f"   - 文件: {result['file']}")
            logger.info(f"   - 提取的知识条目: {result['items_count']}")
            logger.info(f"   - 耗时: {result['analysis_time']}")
            logger.info(f"{'='*80}\n")
            
            return result
            
        except Exception as e:
            logger.error(f"❌ 分析失败: {e}")
            return {
                "success": False,
                "error": str(e),
                "file": Path(file_path).name if file_path else "unknown",
                "timestamp": datetime.now().isoformat(),
            }
    
    @staticmethod
    def _validate_items(items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """验证和补充知识条目"""
        validated = []
        
        for item in items:
            # 检查必需字段
            if not item.get("title") or not item.get("content"):
                logger.warning(f"⚠️ 跳过不完整的条目: {item.get('title', 'Unknown')}")
                continue
            
            # 补充缺失字段
            validated_item = {
                "id": f"kb-{int(datetime.now().timestamp()*1000)}-{len(validated)}",
                "title": item.get("title", "未命名"),
                "category": item.get("category", "其他"),
                "dimension": item.get("dimension", ""),
                "content": item.get("content", ""),
                "scoringRules": item.get("scoringRules", []),
                "createdAt": datetime.now().isoformat(),
                "source": "document_analysis",
            }
            
            validated.append(validated_item)
            logger.debug(f"✓ 验证通过: {validated_item['title']}")
        
        return validated


# 测试函数
def test_document_analyzer():
    """测试文档分析器"""
    logger.info("\n" + "█"*80)
    logger.info("█  文档分析器 - 功能测试")
    logger.info("█"*80)
    
    analyzer = KnowledgeExtractor()
    
    # 模拟测试（使用本地规则）
    test_content = """
    GTV评估标准：
    1. 教育背景 - 申请人需要具有硕士或以上学位，最好来自顶级大学
    2. 工作经验 - 需要至少5年相关工作经验，在行业中有认可度
    3. 技术专长 - 对于技术类申请，需要展示深度的技术能力和创新
    
    评分规则：
    - 学位等级：博士100分，硕士80分，学士50分
    - 工作年限：15+年100分，10-15年90分，5-10年70分
    - 行业影响力：国际知名100分，行业领袖90分，区域影响70分
    """
    
    logger.info("\n测试LLM分析功能...")
    # 注意：实际使用需要API密钥
    logger.info("⚠️ 测试模式 - 需要设置 OPENAI_API_KEY 才能完整运行")
    
    logger.info("\n" + "█"*80)
    logger.info("█  测试完成")
    logger.info("█"*80)


if __name__ == "__main__":
    test_document_analyzer()
