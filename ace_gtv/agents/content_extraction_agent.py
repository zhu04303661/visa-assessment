#!/usr/bin/env python3
"""
GTV签证文案制作 - 内容提取Agent
从项目材料中智能提取和结构化内容，支持多格式文件处理

重构说明：使用 DAO 层进行所有数据库操作，不再直接使用 sqlite3
"""

import os
import json
from datetime import datetime
from typing import Dict, Any, List, Optional
from pathlib import Path

from utils.logger_config import setup_module_logger
from database.dao import (
    get_extraction_dao, get_classification_dao, get_material_dao
)

logger = setup_module_logger("content_extraction_agent", os.getenv("LOG_LEVEL", "INFO"))


class ContentExtractionAgent:
    """
    内容提取Agent
    
    职责：
    1. 从项目文件中提取文本内容
    2. 对提取的内容进行智能分段和标注
    3. 识别关键信息（成就、数据、引用等）
    4. 维护内容与来源的溯源关系
    """
    
    def __init__(self, db_path: str = None, upload_folder: str = None):
        # 使用环境变量配置的路径，或传入的参数
        self.db_path = db_path or os.getenv("COPYWRITING_DB_PATH", "copywriting.db")
        self.upload_folder = upload_folder or os.getenv("UPLOAD_FOLDER", "./uploads")
        self.llm_client = None
        
        # 初始化 DAO 层
        self.extraction_dao = get_extraction_dao()
        self.classification_dao = get_classification_dao()
        self.material_dao = get_material_dao()
        
        self._init_llm()
        logger.info("内容提取Agent初始化完成")
    
    def _init_llm(self):
        """初始化LLM客户端"""
        try:
            from openai import OpenAI
            
            api_key = os.getenv("ENNCLOUD_API_KEY") or os.getenv("OPENAI_API_KEY")
            base_url = os.getenv("ENNCLOUD_BASE_URL") or os.getenv("OPENAI_BASE_URL")
            
            if api_key and base_url:
                self.llm_client = OpenAI(api_key=api_key, base_url=base_url)
                self.model = os.getenv("ENNCLOUD_MODEL", "glm-4.6-no-think")
                logger.info("LLM客户端初始化成功")
            else:
                logger.warning("未配置LLM，将使用规则提取模式")
        except Exception as e:
            logger.error(f"初始化LLM失败: {e}")
    
    def _log_extraction(self, project_id: str, log_type: str, action: str, 
                        file_name: str = None, prompt: str = None, 
                        response: str = None, status: str = 'success', 
                        error_message: str = None):
        """记录提取过程日志（使用 DAO 层）"""
        try:
            self.extraction_dao.add_log(
                project_id=project_id,
                file_name=file_name,
                extraction_type=log_type,
                status=status,
                content=response,
                error_message=error_message
            )
        except Exception as e:
            logger.warning(f"记录日志失败: {e}")
    
    def get_extraction_logs(self, project_id: str) -> List[Dict]:
        """获取项目的提取日志（使用 DAO 层）"""
        try:
            logs = self.extraction_dao.get_logs(project_id, limit=100)
            # 截断响应并转换字段名
            for log in logs:
                if log.get('content'):
                    log['response'] = log['content'][:500]
                log['log_type'] = log.get('extraction_type')
            return logs
        except Exception as e:
            logger.error(f"获取提取日志失败: {e}")
            return []
    
    def extract_project_files(self, project_id: str) -> Dict[str, Any]:
        """
        提取项目所有文件内容
        
        Args:
            project_id: 项目ID
            
        Returns:
            提取结果统计
        """
        try:
            # 记录开始提取
            self._log_extraction(
                project_id, 'extraction', '开始提取项目文件',
                status='started'
            )
            
            # 获取项目文件列表
            files = self._get_project_files(project_id)
            
            # 记录文件列表
            file_list = [f.get('filename', 'unknown') for f in files]
            self._log_extraction(
                project_id, 'extraction', '获取文件列表',
                response=f"找到 {len(files)} 个文件: {', '.join(file_list)}"
            )
            
            if not files:
                self._log_extraction(
                    project_id, 'extraction', '提取完成',
                    response='没有找到需要提取的文件'
                )
                return {
                    "success": True,
                    "data": {
                        "total_files": 0,
                        "processed_files": 0,
                        "total_content_blocks": 0
                    }
                }
            
            total_blocks = 0
            processed_files = 0
            
            for file_info in files:
                try:
                    self._log_extraction(
                        project_id, 'file_processing', f'开始处理文件',
                        file_name=file_info.get('filename'),
                        status='processing'
                    )
                    
                    blocks = self._extract_file_content(project_id, file_info)
                    if blocks:
                        total_blocks += len(blocks)
                        processed_files += 1
                        self._log_extraction(
                            project_id, 'file_processing', f'文件处理完成',
                            file_name=file_info.get('filename'),
                            response=f'提取了 {len(blocks)} 个内容块'
                        )
                except Exception as e:
                    logger.error(f"处理文件失败 {file_info.get('filename')}: {e}")
                    self._log_extraction(
                        project_id, 'file_processing', f'文件处理失败',
                        file_name=file_info.get('filename'),
                        status='error',
                        error_message=str(e)
                    )
            
            # 如果有LLM，检测英文内容并翻译成中文（只保留中文，不保留英文原文）
            if self.llm_client and total_blocks > 0:
                self._translate_english_content(project_id)
            
            # 清理重复内容（减少上下文长度）
            dedup_result = self.deduplicate_content(project_id)
            duplicates_removed = dedup_result.get("data", {}).get("removed", 0) if dedup_result.get("success") else 0
            
            # 生成内容大纲（汇总所有材料的关键信息）
            outline_result = None
            if self.llm_client and total_blocks > 0:
                outline_result = self.generate_content_outline(project_id)
            
            # 记录完成
            self._log_extraction(
                project_id, 'extraction', '提取完成',
                response=f'处理 {processed_files}/{len(files)} 个文件，共 {total_blocks} 个内容块，去除 {duplicates_removed} 个重复'
            )
            
            return {
                "success": True,
                "data": {
                    "total_files": len(files),
                    "processed_files": processed_files,
                    "total_content_blocks": total_blocks,
                    "duplicates_removed": duplicates_removed,
                    "outline_generated": outline_result.get("success", False) if outline_result else False
                }
            }
            
        except Exception as e:
            logger.error(f"提取项目文件失败: {e}")
            return {"success": False, "error": str(e)}
    
    def _get_project_files(self, project_id: str) -> List[Dict]:
        """获取项目的所有文件（使用 DAO 层）"""
        try:
            files = []
            seen_ids = set()
            
            rows = self.material_dao.get_files_by_project(project_id)
            for row in rows:
                file_id = row['id']
                if file_id not in seen_ids:
                    seen_ids.add(file_id)
                    files.append({
                        "id": file_id,
                        "filename": row['file_name'],
                        "file_path": row.get('file_path'),
                        "file_type": row.get('file_type'),
                        "category": row.get('category_id'),
                        "item_id": row.get('item_id'),
                        "description": row.get('description'),
                        "source_table": "material_files",
                        "storage_type": row.get('storage_type') or 'local',
                        "object_bucket": row.get('object_bucket'),
                        "object_key": row.get('object_key')
                    })
            
            logger.info(f"项目 {project_id} 共找到 {len(files)} 个文件待提取")
            return files
            
        except Exception as e:
            logger.error(f"获取项目文件列表失败: {e}")
            return []
    
    def _extract_file_content(self, project_id: str, file_info: Dict) -> List[Dict]:
        """
        从单个文件提取内容
        
        支持的文件类型:
        - PDF: 按页提取
        - Word (docx): 按段落提取
        - Excel (xlsx): 按工作表提取
        - Text/Markdown: 按段落提取
        - JSON: 结构化提取
        
        支持的存储类型:
        - local: 本地文件系统
        - minio/其他: 对象存储（使用内存流）
        
        使用统一文件存储接口，自动处理任何存储后端
        """
        file_type = file_info.get('file_type', '').lower()
        filename = file_info.get('filename', '')
        
        # 如果file_type是通用类型（如document），从文件扩展名推断实际类型
        if file_type in ['document', 'spreadsheet', 'image', 'other', '']:
            ext = Path(filename).suffix.lower().lstrip('.')
            if ext:
                file_type = ext
        file_id = file_info.get('id')
        
        try:
            # 使用统一接口获取文件内容（自动处理任何存储类型）
            file_bytes = self._get_file_bytes(file_info, filename)
            if file_bytes is None:
                return []
            
            # 使用内存流提取内容
            content_blocks = self._extract_from_bytes(file_bytes, file_type, filename, file_id)
            
            # 保存到数据库
            if content_blocks:
                self._save_content_blocks(project_id, content_blocks)
            
            return content_blocks
            
        except Exception as e:
            logger.error(f"提取文件内容失败 {filename}: {e}")
            return []
    
    def _get_file_bytes(self, file_info: Dict, filename: str) -> bytes:
        """
        从存储系统获取文件内容到内存
        
        使用统一文件存储接口，自动处理 MinIO 和本地存储
        """
        try:
            from database.file_storage import get_file_from_db_record
            
            content = get_file_from_db_record(file_info)
            if content:
                logger.info(f"文件已加载到内存: {filename} ({len(content)} bytes)")
                return content
            
            logger.warning(f"无法获取文件: {filename}")
            return None
            
        except Exception as e:
            logger.warning(f"获取文件失败: {filename}, {e}")
            return None
    
    def _resolve_local_path(self, file_path: str) -> str:
        """解析本地文件路径"""
        if not file_path:
            return None
            
        if os.path.isabs(file_path):
            return file_path
        
        # 清理路径前缀
        cleaned_path = file_path
        
        # 移除开头的 ./ 
        if cleaned_path.startswith('./'):
            cleaned_path = cleaned_path[2:]
        
        # 检查路径是否已经包含uploads前缀
        if cleaned_path.startswith('uploads/') or cleaned_path.startswith('uploads\\'):
            return cleaned_path
        else:
            return os.path.join(self.upload_folder, cleaned_path)
    
    def _extract_from_bytes(self, file_bytes: bytes, file_type: str, filename: str, file_id: int) -> List[Dict]:
        """从内存字节流提取内容（用于 MinIO 文件）"""
        from io import BytesIO
        
        content_blocks = []
        
        try:
            if file_type == 'pdf':
                content_blocks = self._extract_pdf_from_bytes(file_bytes, filename, file_id)
            elif file_type in ['docx', 'doc']:
                content_blocks = self._extract_docx_from_bytes(file_bytes, filename, file_id)
            elif file_type in ['xlsx', 'xls']:
                content_blocks = self._extract_excel_from_bytes(file_bytes, filename, file_id)
            elif file_type in ['txt', 'md', 'json']:
                content_blocks = self._extract_text_from_bytes(file_bytes, filename, file_id)
            else:
                logger.info(f"跳过不支持的文件类型: {file_type}")
        except Exception as e:
            logger.error(f"从内存提取文件内容失败 {filename}: {e}")
        
        return content_blocks
    
    def _extract_from_file(self, file_path: str, file_type: str, filename: str, file_id: int) -> List[Dict]:
        """从本地文件提取内容"""
        content_blocks = []
        
        try:
            if file_type == 'pdf':
                content_blocks = self._extract_pdf(file_path, filename, file_id)
            elif file_type in ['docx', 'doc']:
                content_blocks = self._extract_docx(file_path, filename, file_id)
            elif file_type in ['xlsx', 'xls']:
                content_blocks = self._extract_excel(file_path, filename, file_id)
            elif file_type in ['txt', 'md', 'json']:
                content_blocks = self._extract_text(file_path, filename, file_id)
            else:
                logger.info(f"跳过不支持的文件类型: {file_type}")
        except Exception as e:
            logger.error(f"从文件提取内容失败 {filename}: {e}")
        
        return content_blocks
    
    def _extract_pdf_from_bytes(self, file_bytes: bytes, filename: str, file_id: int) -> List[Dict]:
        """从内存字节流提取 PDF 内容"""
        blocks = []
        
        try:
            import fitz  # PyMuPDF
            
            # 从内存打开 PDF
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                
                if text.strip():
                    blocks.append({
                        "source_file": filename,
                        "source_file_id": file_id,
                        "block_type": "page",
                        "block_index": page_num + 1,
                        "content": text.strip(),
                        "metadata": {
                            "page_number": page_num + 1,
                            "total_pages": len(doc)
                        }
                    })
            
            doc.close()
            logger.info(f"PDF 提取完成: {filename}, {len(blocks)} 个内容块")
            
        except Exception as e:
            logger.error(f"PDF 内存提取失败 {filename}: {e}")
        
        return blocks
    
    def _extract_docx_from_bytes(self, file_bytes: bytes, filename: str, file_id: int) -> List[Dict]:
        """从内存字节流提取 Word 文档内容"""
        from io import BytesIO
        blocks = []
        
        try:
            from docx import Document
            
            # 从内存打开 Word 文档
            doc = Document(BytesIO(file_bytes))
            
            block_index = 0
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    block_index += 1
                    blocks.append({
                        "source_file": filename,
                        "source_file_id": file_id,
                        "block_type": "paragraph",
                        "block_index": block_index,
                        "content": text,
                        "metadata": {
                            "style": para.style.name if para.style else None
                        }
                    })
            
            # 提取表格内容
            for table_idx, table in enumerate(doc.tables):
                table_content = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_content.append(" | ".join(row_data))
                
                if table_content:
                    block_index += 1
                    blocks.append({
                        "source_file": filename,
                        "source_file_id": file_id,
                        "block_type": "table",
                        "block_index": block_index,
                        "content": "\n".join(table_content),
                        "metadata": {
                            "table_index": table_idx + 1
                        }
                    })
            
            logger.info(f"Word 提取完成: {filename}, {len(blocks)} 个内容块")
            
        except Exception as e:
            logger.error(f"Word 内存提取失败 {filename}: {e}")
        
        return blocks
    
    def _extract_excel_from_bytes(self, file_bytes: bytes, filename: str, file_id: int) -> List[Dict]:
        """从内存字节流提取 Excel 内容"""
        from io import BytesIO
        blocks = []
        
        try:
            import openpyxl
            
            # 从内存打开 Excel
            wb = openpyxl.load_workbook(BytesIO(file_bytes), data_only=True)
            
            block_index = 0
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                rows_content = []
                
                for row in sheet.iter_rows(values_only=True):
                    # 过滤空行
                    row_values = [str(cell) if cell is not None else '' for cell in row]
                    if any(v.strip() for v in row_values):
                        rows_content.append(" | ".join(row_values))
                
                if rows_content:
                    block_index += 1
                    blocks.append({
                        "source_file": filename,
                        "source_file_id": file_id,
                        "block_type": "sheet",
                        "block_index": block_index,
                        "content": "\n".join(rows_content),
                        "metadata": {
                            "sheet_name": sheet_name,
                            "row_count": len(rows_content)
                        }
                    })
            
            wb.close()
            logger.info(f"Excel 提取完成: {filename}, {len(blocks)} 个内容块")
            
        except Exception as e:
            logger.error(f"Excel 内存提取失败 {filename}: {e}")
        
        return blocks
    
    def _extract_text_from_bytes(self, file_bytes: bytes, filename: str, file_id: int) -> List[Dict]:
        """从内存字节流提取文本内容"""
        blocks = []
        
        try:
            # 尝试多种编码
            content = None
            for encoding in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
                try:
                    content = file_bytes.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                logger.warning(f"无法解码文本文件: {filename}")
                return []
            
            # 按段落分割
            paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
            
            for idx, para in enumerate(paragraphs):
                if para:
                    blocks.append({
                        "source_file": filename,
                        "source_file_id": file_id,
                        "block_type": "paragraph",
                        "block_index": idx + 1,
                        "content": para,
                        "metadata": {}
                    })
            
            logger.info(f"文本提取完成: {filename}, {len(blocks)} 个内容块")
            
        except Exception as e:
            logger.error(f"文本内存提取失败 {filename}: {e}")
        
        return blocks
    
    def _extract_pdf(self, file_path: str, filename: str, file_id: int) -> List[Dict]:
        """提取PDF内容"""
        blocks = []
        
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(file_path)
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                
                if text.strip():
                    blocks.append({
                        "file_id": file_id,
                        "source_file": filename,
                        "source_page": page_num + 1,
                        "content_type": "text",
                        "content": text.strip(),
                        "word_count": len(text.split())
                    })
            
            doc.close()
            
        except ImportError:
            logger.warning("PyMuPDF未安装，尝试使用备用方案")
            try:
                import PyPDF2
                
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    for page_num, page in enumerate(reader.pages):
                        text = page.extract_text()
                        if text.strip():
                            blocks.append({
                                "file_id": file_id,
                                "source_file": filename,
                                "source_page": page_num + 1,
                                "content_type": "text",
                                "content": text.strip(),
                                "word_count": len(text.split())
                            })
            except Exception as e:
                logger.error(f"PDF提取失败: {e}")
        
        return blocks
    
    def _extract_docx(self, file_path: str, filename: str, file_id: int) -> List[Dict]:
        """提取Word文档内容"""
        blocks = []
        
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # 合并段落为有意义的块
            current_block = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    current_block.append(text)
                elif current_block:
                    # 空段落作为分隔
                    full_text = "\n".join(current_block)
                    if len(full_text) > 50:  # 只保留有意义的块
                        blocks.append({
                            "file_id": file_id,
                            "source_file": filename,
                            "source_page": None,
                            "content_type": "text",
                            "content": full_text,
                            "word_count": len(full_text.split())
                        })
                    current_block = []
            
            # 处理最后一个块
            if current_block:
                full_text = "\n".join(current_block)
                if len(full_text) > 50:
                    blocks.append({
                        "file_id": file_id,
                        "source_file": filename,
                        "source_page": None,
                        "content_type": "text",
                        "content": full_text,
                        "word_count": len(full_text.split())
                    })
            
            # 提取表格
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(" | ".join(row_data))
                
                if table_data:
                    table_text = "\n".join(table_data)
                    blocks.append({
                        "file_id": file_id,
                        "source_file": filename,
                        "source_page": None,
                        "content_type": "table",
                        "content": table_text,
                        "word_count": len(table_text.split())
                    })
            
        except Exception as e:
            logger.error(f"Word文档提取失败: {e}")
        
        return blocks
    
    def _extract_excel(self, file_path: str, filename: str, file_id: int) -> List[Dict]:
        """提取Excel内容"""
        blocks = []
        
        try:
            import pandas as pd
            
            xl = pd.ExcelFile(file_path)
            
            for sheet_name in xl.sheet_names:
                df = pd.read_excel(xl, sheet_name=sheet_name)
                
                # 转换为文本
                text_content = df.to_string(index=False)
                
                if len(text_content.strip()) > 50:
                    blocks.append({
                        "file_id": file_id,
                        "source_file": f"{filename} - {sheet_name}",
                        "source_page": None,
                        "content_type": "table",
                        "content": text_content,
                        "word_count": len(text_content.split())
                    })
            
        except Exception as e:
            logger.error(f"Excel提取失败: {e}")
        
        return blocks
    
    def _extract_text(self, file_path: str, filename: str, file_id: int) -> List[Dict]:
        """提取文本文件内容"""
        blocks = []
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            if len(content.strip()) > 50:
                blocks.append({
                    "file_id": file_id,
                    "source_file": filename,
                    "source_page": None,
                    "content_type": "text",
                    "content": content.strip(),
                    "word_count": len(content.split())
                })
            
        except Exception as e:
            logger.error(f"文本文件提取失败: {e}")
        
        return blocks
    
    def _save_content_blocks(self, project_id: str, blocks: List[Dict]):
        """保存内容块到数据库（使用 DAO 层，自动去除重复内容）"""
        try:
            # 获取已存在的内容哈希，用于去重
            existing_texts = self.extraction_dao.get_all_content_texts(project_id)
            existing_hashes = set()
            for text in existing_texts:
                content_hash = self._get_content_hash(text)
                existing_hashes.add(content_hash)
            
            saved_count = 0
            skipped_count = 0
            
            for block in blocks:
                content = block.get("content", "")
                
                # 计算内容哈希，检查是否重复
                content_hash = self._get_content_hash(content)
                if content_hash in existing_hashes:
                    skipped_count += 1
                    continue
                
                # 标记为已存在，避免批次内重复
                existing_hashes.add(content_hash)
                
                self.extraction_dao.save_content_block(
                    project_id=project_id,
                    file_id=block.get("file_id"),
                    source_file=block.get("source_file"),
                    source_page=block.get("source_page"),
                    content_type=block.get("content_type", "text"),
                    content=content,
                    word_count=block.get("word_count", 0)
                )
                saved_count += 1
            
            if skipped_count > 0:
                logger.info(f"保存 {saved_count} 个内容块到项目 {project_id}，跳过 {skipped_count} 个重复内容")
            else:
                logger.info(f"保存 {saved_count} 个内容块到项目 {project_id}")
            
        except Exception as e:
            logger.error(f"保存内容块失败: {e}")
    
    def _get_content_hash(self, content: str) -> str:
        """
        计算内容的标准化哈希值，用于去重
        
        标准化处理：
        - 去除首尾空白
        - 将多个空白字符合并为单个空格
        - 转为小写进行比较
        """
        import hashlib
        import re
        
        if not content:
            return ""
        
        # 标准化处理
        normalized = content.strip().lower()
        normalized = re.sub(r'\s+', ' ', normalized)  # 合并空白字符
        
        # 计算MD5哈希
        return hashlib.md5(normalized.encode('utf-8')).hexdigest()
    
    def deduplicate_content(self, project_id: str) -> Dict[str, Any]:
        """
        清理项目中的重复内容（使用 DAO 层）
        
        检查并删除：
        1. 完全相同的内容（哈希相同）
        2. 高度相似的内容（相似度>90%）
        """
        try:
            # 获取所有内容块
            rows = self.extraction_dao.get_contents_for_dedup(project_id)
            
            if not rows:
                return {"success": True, "data": {"removed": 0, "message": "没有内容可清理"}}
            
            # 检测重复内容
            seen_hashes = {}  # hash -> (id, source_file)
            duplicates = []
            
            for row in rows:
                content_hash = self._get_content_hash(row.get('content', ''))
                
                if content_hash in seen_hashes:
                    duplicates.append({
                        "id": row['id'],
                        "source_file": row.get('source_file'),
                        "duplicate_of": seen_hashes[content_hash]
                    })
                else:
                    seen_hashes[content_hash] = (row['id'], row.get('source_file'))
            
            # 删除重复内容
            removed_count = 0
            for dup in duplicates:
                try:
                    self.extraction_dao.delete_content_by_id(dup['id'])
                    removed_count += 1
                except Exception as e:
                    logger.warning(f"删除重复内容失败 (id={dup['id']}): {e}")
            
            if removed_count > 0:
                self._log_extraction(
                    project_id, 'deduplication', '清理重复内容',
                    response=f'删除了 {removed_count} 个重复内容块'
                )
            
            logger.info(f"项目 {project_id} 去重完成，删除 {removed_count} 个重复内容块")
            
            return {
                "success": True,
                "data": {
                    "total_checked": len(rows),
                    "duplicates_found": len(duplicates),
                    "removed": removed_count
                }
            }
            
        except Exception as e:
            logger.error(f"去重失败: {e}")
            return {"success": False, "error": str(e)}
    
    def _translate_english_content(self, project_id: str):
        """检测英文内容并翻译为中文（使用 DAO 层）"""
        if not self.llm_client:
            return
        
        try:
            # 获取未处理的内容块
            rows = self.extraction_dao.get_contents_with_source(project_id)
            
            # 过滤未处理的内容（structured_data 为空表示未处理）
            unprocessed_rows = [r for r in rows if not r.get('structured_data')]
            
            if not unprocessed_rows:
                logger.info(f"项目 {project_id} 没有需要处理的内容块")
                return
            
            # 统计英文内容块
            english_blocks = []
            for row in unprocessed_rows:
                content = row.get('content', '')
                if self._is_english_content(content):
                    english_blocks.append(row)
            
            if not english_blocks:
                logger.info(f"项目 {project_id} 没有英文内容需要翻译")
                # 标记所有内容为已处理
                for row in unprocessed_rows:
                    self.extraction_dao.update_content_metadata(row['id'], '{"processed": true}')
                return
            
            logger.info(f"检测到 {len(english_blocks)} 个英文内容块，开始翻译...")
            self._log_extraction(
                project_id, 'translation', '开始翻译英文内容',
                response=f'检测到 {len(english_blocks)} 个英文内容块'
            )
            
            translated_count = 0
            for i, row in enumerate(english_blocks):
                block_id = row['id']
                content = row.get('content', '')
                source_file = row.get('source_file', '')
                
                if len(content) < 50:
                    continue
                
                try:
                    result, prompt_used = self._translate_to_chinese(content[:4000], source_file, return_prompt=True)
                    
                    if result:
                        if translated_count < 3:
                            self._log_extraction(
                                project_id, 'translation', f'翻译内容块 {block_id}',
                                file_name=source_file,
                                prompt=prompt_used[:2000] if prompt_used else None,
                                response=result.get("chinese_translation", "")[:500]
                            )
                        
                        new_content = result.get("chinese_translation", content)
                        metadata = json.dumps({
                            "translated": True,
                            "original_language": "en",
                            "key_info": result.get("key_info", {})
                        }, ensure_ascii=False)
                        
                        self.extraction_dao.update_content_text(block_id, new_content)
                        self.extraction_dao.update_content_metadata(block_id, metadata)
                        translated_count += 1
                        logger.info(f"翻译进度: {translated_count}/{len(english_blocks)} - 内容块 {block_id}")
                except Exception as e:
                    logger.warning(f"翻译内容块 {block_id} 失败: {e}")
                    self._log_extraction(
                        project_id, 'translation', f'翻译失败',
                        file_name=source_file,
                        status='error',
                        error_message=str(e)
                    )
                    try:
                        self.extraction_dao.update_content_metadata(
                            block_id, '{"processed": true, "translation_failed": true}'
                        )
                    except:
                        pass
                    continue
            
            logger.info(f"完成项目 {project_id} 的英文内容翻译，共翻译 {translated_count}/{len(english_blocks)} 个内容块")
            
        except Exception as e:
            logger.error(f"翻译英文内容失败: {e}")
    
    def _is_english_content(self, content: str) -> bool:
        """检测内容是否主要是英文"""
        if not content or len(content) < 20:
            return False
        
        # 统计字符类型
        english_chars = sum(1 for c in content if c.isascii() and c.isalpha())
        chinese_chars = sum(1 for c in content if '\u4e00' <= c <= '\u9fff')
        total_alpha = english_chars + chinese_chars
        
        if total_alpha == 0:
            return False
        
        # 如果英文字符占比超过70%，认为是英文内容
        return (english_chars / total_alpha) > 0.7
    
    def _translate_to_chinese(self, content: str, source_file: str = "", return_prompt: bool = False):
        """
        将英文内容翻译为中文（只保留中文，不保留英文原文，减少上下文长度）
        
        Args:
            content: 要翻译的内容
            source_file: 来源文件名
            return_prompt: 是否返回使用的提示词
            
        Returns:
            如果return_prompt=True，返回(result, prompt)元组
            否则只返回result
        """
        if not self.llm_client:
            return (None, None) if return_prompt else None
        
        try:
            prompt = f"""请将以下英文内容翻译成中文，并提取关键信息。

来源文件: {source_file}

英文原文:
{content}

请按以下JSON格式返回：
{{
    "chinese_translation": "翻译后的中文内容（精简准确，去除冗余）",
    "key_info": {{
        "summary": "一句话中文摘要（30字以内）",
        "key_points": ["核心要点1", "核心要点2"],
        "achievements": ["成就/奖项（如有）"],
        "numbers": ["关键数据/数字（如有）"],
        "entities": ["重要人名/公司/机构（如有）"]
    }}
}}

翻译要求：
1. 只保留核心信息，去除重复和无意义内容
2. 人名保留英文，公司名/学校名可添加中文注释
3. 数据和指标要保留完整
4. 只返回JSON格式，不要其他文字"""
            
            response = self.llm_client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=2000,
                temperature=0.3
            )
            
            result_text = response.choices[0].message.content
            
            import re
            json_match = re.search(r'\{[\s\S]*\}', result_text)
            if json_match:
                result = json.loads(json_match.group())
                return (result, prompt) if return_prompt else result
            
            return (None, prompt) if return_prompt else None
            
        except Exception as e:
            logger.warning(f"翻译内容失败: {e}")
            return (None, prompt if 'prompt' in locals() else None) if return_prompt else None
    
    def _enhance_extracted_content(self, project_id: str):
        """使用LLM增强提取的内容 - 已整合到 _translate_english_content"""
        pass
    
    def _extract_key_entities(self, content: str) -> Dict:
        """使用LLM提取关键实体"""
        if not self.llm_client:
            return {}
        
        try:
            prompt = f"""
请从以下文本中提取关键信息，以JSON格式返回：

{content}

返回格式：
{{
    "achievements": ["成就1", "成就2"],
    "numbers": ["数据指标1", "数据指标2"],
    "organizations": ["组织名1", "组织名2"],
    "technologies": ["技术1", "技术2"],
    "dates": ["日期1", "日期2"]
}}

只返回JSON，不要其他文字。
"""
            
            response = self.llm_client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=500,
                temperature=0.3
            )
            
            result_text = response.choices[0].message.content
            
            import re
            json_match = re.search(r'\{[\s\S]*\}', result_text)
            if json_match:
                return json.loads(json_match.group())
            
            return {}
            
        except Exception as e:
            logger.warning(f"提取关键实体失败: {e}")
            return {}
    
    def get_project_context(self, project_id: str, with_sources: bool = True) -> Dict[str, Any]:
        """
        获取项目的完整上下文（使用 DAO 层）
        """
        try:
            # 获取内容块
            blocks = self.extraction_dao.get_contents_with_source(project_id)
            
            context_parts = []
            total_words = 0
            
            for block in blocks:
                if with_sources:
                    source = f"\n=== {block.get('source_file', 'unknown')}"
                    source += " ===\n"
                    context_parts.append(source)
                
                content = block.get('content', '')
                context_parts.append(content)
                total_words += len(content.split()) if content else 0
            
            # 获取统计信息
            stats = self.extraction_dao.get_content_stats(project_id)
            
            # 获取 material_files 表中的可用文件数量
            material_files = self.material_dao.get_files_by_project(project_id)
            available_files = len(material_files)
            
            return {
                "success": True,
                "data": {
                    "context": "\n".join(context_parts),
                    "total_files": available_files,
                    "extracted_files": stats.get('file_count', 0),
                    "total_blocks": stats.get('content_count', 0),
                    "total_words": total_words,
                    "last_updated": None
                }
            }
            
        except Exception as e:
            logger.error(f"获取项目上下文失败: {e}")
            return {"success": False, "error": str(e)}
    
    def get_content_blocks(self, project_id: str, file_id: int = None, 
                          content_type: str = None) -> Dict[str, Any]:
        """获取项目内容块列表（使用 DAO 层）"""
        try:
            blocks = self.extraction_dao.get_contents_with_source(project_id, content_type=content_type)
            
            # 过滤文件ID
            if file_id:
                blocks = [b for b in blocks if b.get('file_id') == file_id]
            
            # 处理内容预览和 metadata
            result_blocks = []
            for block in blocks:
                processed = {
                    'id': block.get('id'),
                    'file_id': block.get('file_id'),
                    'source_file': block.get('source_file'),
                    'content_type': block.get('content_type'),
                    'content_preview': (block.get('content') or '')[:200]
                }
                result_blocks.append(processed)
            
            return {
                "success": True,
                "data": result_blocks,
                "total": len(result_blocks)
            }
            
        except Exception as e:
            logger.error(f"获取内容块失败: {e}")
            return {"success": False, "error": str(e)}
    
    def search_content(self, project_id: str, keyword: str) -> Dict[str, Any]:
        """搜索项目内容（使用 DAO 层）"""
        try:
            rows = self.extraction_dao.search_content(project_id, keyword)
            
            results = []
            for row in rows:
                content = row.get('content', '')
                # 找到关键词位置，截取上下文
                pos = content.lower().find(keyword.lower())
                start = max(0, pos - 100)
                end = min(len(content), pos + len(keyword) + 100)
                
                snippet = content[start:end]
                if start > 0:
                    snippet = "..." + snippet
                if end < len(content):
                    snippet = snippet + "..."
                
                results.append({
                    "id": row['id'],
                    "source_file": row.get('source_file'),
                    "content_type": row.get('content_type'),
                    "snippet": snippet
                })
            
            return {
                "success": True,
                "data": results,
                "total": len(results),
                "keyword": keyword
            }
            
        except Exception as e:
            logger.error(f"搜索内容失败: {e}")
            return {"success": False, "error": str(e)}
    
    def generate_content_outline(self, project_id: str) -> Dict[str, Any]:
        """
        生成内容大纲 - 汇总所有材料的关键信息，提炼关键词和脉络
        
        大纲结构：
        - 文件清单：每个文件的主要内容摘要
        - 关键词云：所有材料的核心关键词
        - 信息脉络：申请人的主要经历、成就、证据线索
        - 材料评估：各类证据的覆盖情况
        
        Args:
            project_id: 项目ID
            
        Returns:
            生成的大纲结果
        """
        try:
            # 获取所有提取的内容（使用 DAO 层）
            rows = self.extraction_dao.get_contents_with_source(project_id)
            
            if not rows:
                return {"success": False, "error": "没有提取的内容可生成大纲"}
            
            # 按文件分组内容
            file_contents = {}
            for row in rows:
                source_file = row.get('source_file') or "未知文件"
                if source_file not in file_contents:
                    file_contents[source_file] = {
                        "texts": [],
                        "word_count": 0,
                        "metadata_list": []
                    }
                content = row.get('content') or ''
                file_contents[source_file]["texts"].append(content)
                file_contents[source_file]["word_count"] += len(content.split())
            
            self._log_extraction(
                project_id, 'outline', '开始生成内容大纲',
                response=f'共 {len(file_contents)} 个文件待处理'
            )
            
            # 使用LLM生成大纲
            if self.llm_client:
                outline = self._ai_generate_outline(project_id, file_contents)
            else:
                outline = self._rule_based_outline(file_contents)
            
            # 保存大纲
            self._save_outline(project_id, outline)
            
            self._log_extraction(
                project_id, 'outline', '内容大纲生成完成',
                response=f'共 {len(outline.get("file_summaries", []))} 个文件，{len(outline.get("keywords", []))} 个关键词'
            )
            
            return {
                "success": True,
                "data": outline
            }
            
        except Exception as e:
            logger.error(f"生成内容大纲失败: {e}")
            return {"success": False, "error": str(e)}
    
    def _ai_generate_outline(self, project_id: str, file_contents: Dict) -> Dict:
        """使用AI生成内容大纲"""
        # 准备文件摘要列表
        file_summaries = []
        all_key_info = []
        
        for filename, data in file_contents.items():
            # 合并该文件的所有内容（限制长度）
            combined_text = "\n".join(data["texts"])[:3000]
            
            # 收集该文件已提取的关键信息
            for meta in data["metadata_list"]:
                if meta.get("key_info"):
                    all_key_info.append({
                        "file": filename,
                        "info": meta["key_info"]
                    })
            
            # 为每个文件生成简短摘要
            try:
                summary = self._summarize_file_content(filename, combined_text)
                if summary:
                    file_summaries.append({
                        "filename": filename,
                        "word_count": data["word_count"],
                        "summary": summary.get("summary", ""),
                        "content_type": summary.get("content_type", "其他"),
                        "key_points": summary.get("key_points", []),
                        "relevance": summary.get("relevance", "一般")
                    })
            except Exception as e:
                logger.warning(f"生成文件摘要失败 {filename}: {e}")
                file_summaries.append({
                    "filename": filename,
                    "word_count": data["word_count"],
                    "summary": "摘要生成失败",
                    "content_type": "未知",
                    "key_points": [],
                    "relevance": "未知"
                })
        
        # 生成整体大纲
        outline = self._generate_overall_outline(project_id, file_summaries, all_key_info)
        
        return outline
    
    def _summarize_file_content(self, filename: str, content: str) -> Optional[Dict]:
        """为单个文件生成内容摘要"""
        if not self.llm_client:
            return None
        
        try:
            prompt = f"""请分析以下文件内容，生成简短摘要。

文件名: {filename}

内容:
{content[:2500]}

请按以下JSON格式返回（控制总字数在200字以内）：
{{
    "summary": "文件主要内容的一句话摘要（30字以内）",
    "content_type": "文件类型（简历/推荐信/专利/论文/证书/媒体报道/其他）",
    "key_points": ["关键要点1", "关键要点2", "关键要点3"],
    "relevance": "与GTV申请的相关度（高/中/低）",
    "entities": ["提到的重要人名/公司/机构"]
}}

只返回JSON，不要其他文字。"""
            
            response = self.llm_client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=500,
                temperature=0.3
            )
            
            result_text = response.choices[0].message.content
            
            import re
            json_match = re.search(r'\{[\s\S]*\}', result_text)
            if json_match:
                return json.loads(json_match.group())
            
            return None
            
        except Exception as e:
            logger.warning(f"文件摘要生成失败 {filename}: {e}")
            return None
    
    def _generate_overall_outline(self, project_id: str, file_summaries: List[Dict], 
                                   all_key_info: List[Dict]) -> Dict:
        """生成整体内容大纲"""
        if not self.llm_client:
            return self._rule_based_outline_from_summaries(file_summaries)
        
        try:
            # 准备输入数据
            summaries_text = json.dumps(file_summaries, ensure_ascii=False, indent=2)
            key_info_text = json.dumps(all_key_info[:20], ensure_ascii=False, indent=2)  # 限制数量
            
            prompt = f"""根据以下材料摘要，生成一个GTV申请的内容大纲。

文件摘要列表:
{summaries_text[:4000]}

已提取的关键信息:
{key_info_text[:2000]}

请生成结构化大纲，按以下JSON格式返回：
{{
    "applicant_profile": {{
        "name": "申请人姓名（如能确定）",
        "current_position": "当前职位",
        "domain": "所属领域",
        "experience_years": "工作年限（数字或范围）"
    }},
    "file_summaries": [
        {{
            "filename": "文件名",
            "type": "类型",
            "summary": "摘要",
            "relevance": "相关度"
        }}
    ],
    "keywords": ["关键词1", "关键词2", "关键词3"],
    "career_timeline": [
        {{
            "period": "时间段",
            "event": "事件描述"
        }}
    ],
    "achievement_categories": {{
        "leadership": ["领导力相关成就"],
        "innovation": ["创新相关成就"],
        "recognition": ["行业认可"],
        "academic": ["学术成就"]
    }},
    "evidence_coverage": {{
        "MC标准覆盖": {{
            "MC1_产品团队领导力": "有/部分/无",
            "MC2_商业发展": "有/部分/无",
            "MC3_非营利组织": "有/部分/无",
            "MC4_专家评审": "有/部分/无"
        }},
        "OC标准覆盖": {{
            "OC1_创新": "有/部分/无",
            "OC2_行业认可": "有/部分/无",
            "OC3_重大贡献": "有/部分/无",
            "OC4_学术贡献": "有/部分/无"
        }}
    }},
    "material_gaps": ["缺少的材料类型1", "缺少的材料类型2"],
    "overall_assessment": "材料整体情况的一句话评价"
}}

只返回JSON，不要其他文字。"""
            
            response = self.llm_client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=2000,
                temperature=0.3
            )
            
            result_text = response.choices[0].message.content
            
            # 记录日志
            self._log_extraction(
                project_id, 'outline', '生成整体大纲',
                prompt=prompt[:3000],
                response=result_text[:2000]
            )
            
            import re
            json_match = re.search(r'\{[\s\S]*\}', result_text)
            if json_match:
                outline = json.loads(json_match.group())
                outline["generated_at"] = datetime.now().isoformat()
                outline["total_files"] = len(file_summaries)
                outline["ai_generated"] = True
                return outline
            
            return self._rule_based_outline_from_summaries(file_summaries)
            
        except Exception as e:
            logger.error(f"生成整体大纲失败: {e}")
            return self._rule_based_outline_from_summaries(file_summaries)
    
    def _rule_based_outline(self, file_contents: Dict) -> Dict:
        """基于规则生成内容大纲（无LLM时的回退方案）"""
        file_summaries = []
        all_keywords = set()
        
        for filename, data in file_contents.items():
            combined_text = " ".join(data["texts"])
            
            # 简单关键词提取
            keywords = self._extract_simple_keywords(combined_text)
            all_keywords.update(keywords)
            
            # 判断文件类型
            content_type = self._guess_file_type(filename, combined_text)
            
            file_summaries.append({
                "filename": filename,
                "word_count": data["word_count"],
                "summary": f"{content_type}文件，约{data['word_count']}字",
                "content_type": content_type,
                "key_points": keywords[:3],
                "relevance": "待评估"
            })
        
        return {
            "file_summaries": file_summaries,
            "keywords": list(all_keywords)[:30],
            "applicant_profile": {},
            "career_timeline": [],
            "achievement_categories": {},
            "evidence_coverage": {},
            "material_gaps": [],
            "overall_assessment": f"共收集 {len(file_summaries)} 个文件，建议使用AI进行深度分析",
            "generated_at": datetime.now().isoformat(),
            "total_files": len(file_summaries),
            "ai_generated": False
        }
    
    def _rule_based_outline_from_summaries(self, file_summaries: List[Dict]) -> Dict:
        """从已有摘要生成大纲（无LLM时的回退方案）"""
        all_keywords = []
        for fs in file_summaries:
            all_keywords.extend(fs.get("key_points", []))
        
        return {
            "file_summaries": file_summaries,
            "keywords": list(set(all_keywords))[:30],
            "applicant_profile": {},
            "career_timeline": [],
            "achievement_categories": {},
            "evidence_coverage": {},
            "material_gaps": [],
            "overall_assessment": f"共收集 {len(file_summaries)} 个文件",
            "generated_at": datetime.now().isoformat(),
            "total_files": len(file_summaries),
            "ai_generated": False
        }
    
    def _extract_simple_keywords(self, text: str) -> List[str]:
        """简单关键词提取"""
        keywords = []
        
        # GTV相关关键词
        gtv_keywords = [
            "领导", "团队", "产品", "创新", "专利", "论文", "发表", 
            "奖项", "荣誉", "认可", "贡献", "影响", "用户", "收入",
            "融资", "投资", "创业", "CTO", "CEO", "创始人", "总监",
            "AI", "人工智能", "机器学习", "深度学习", "数据", "算法",
            "区块链", "金融科技", "数字健康", "云计算", "大数据"
        ]
        
        text_lower = text.lower()
        for kw in gtv_keywords:
            if kw.lower() in text_lower:
                keywords.append(kw)
        
        return keywords[:10]
    
    def _guess_file_type(self, filename: str, content: str) -> str:
        """猜测文件类型"""
        filename_lower = filename.lower()
        content_lower = content.lower()
        
        if "cv" in filename_lower or "resume" in filename_lower or "简历" in filename:
            return "简历"
        elif "recommendation" in filename_lower or "推荐" in filename:
            return "推荐信"
        elif "patent" in filename_lower or "专利" in filename:
            return "专利"
        elif "paper" in filename_lower or "论文" in filename:
            return "论文"
        elif "certificate" in filename_lower or "证书" in filename:
            return "证书"
        elif "award" in content_lower or "奖" in content:
            return "奖项证明"
        else:
            return "其他"
    
    def _save_outline(self, project_id: str, outline: Dict):
        """保存内容大纲到数据库（使用 DAO 层）"""
        try:
            self.extraction_dao.save_outline(
                project_id=project_id,
                outline_data=outline,
                total_files=outline.get("total_files", 0),
                total_keywords=len(outline.get("keywords", []))
            )
            logger.info(f"保存内容大纲: {project_id}")
        except Exception as e:
            logger.error(f"保存内容大纲失败: {e}")
    
    def get_content_outline(self, project_id: str) -> Dict[str, Any]:
        """获取项目的内容大纲（使用 DAO 层）"""
        try:
            row = self.extraction_dao.get_outline(project_id)
            
            if row:
                return {
                    "success": True,
                    "data": {
                        "outline": json.loads(row['outline_data']) if row.get('outline_data') else {},
                        "total_files": row.get('total_files', 0),
                        "total_keywords": row.get('total_keywords', 0),
                        "updated_at": row.get('updated_at')
                    }
                }
            
            return {"success": True, "data": None}
            
        except Exception as e:
            logger.error(f"获取内容大纲失败: {e}")
            return {"success": False, "error": str(e)}

    def clear_project_content(self, project_id: str) -> Dict[str, Any]:
        """清除项目的提取内容（使用 DAO 层）"""
        try:
            deleted = self.extraction_dao.delete_contents(project_id)
            
            return {
                "success": True,
                "data": {"deleted_blocks": deleted}
            }
            
        except Exception as e:
            logger.error(f"清除项目内容失败: {e}")
            return {"success": False, "error": str(e)}

    def clear_all_extraction_data(self, project_id: str) -> Dict[str, Any]:
        """
        清除项目的所有提取数据（使用 DAO 层）
        """
        try:
            stats = {}
            
            # 清除提取的内容块
            stats['deleted_contents'] = self.extraction_dao.delete_contents(project_id)
            
            # 清除提取日志
            stats['deleted_logs'] = self.extraction_dao.delete_logs(project_id)
            
            # 清除内容分类
            stats['deleted_classifications'] = self.classification_dao.delete_classifications(project_id)
            
            logger.info(f"项目 {project_id} 提取数据已清理: {stats}")
            
            return {
                "success": True,
                "data": {
                    "message": "提取数据已清理完成，可以重新提取",
                    "stats": stats
                }
            }
            
        except Exception as e:
            logger.error(f"清除项目提取数据失败: {e}")
            return {"success": False, "error": str(e)}

    # ==================== 内容分类功能 ====================
    
    # 分类标准定义
    CLASSIFICATION_CATEGORIES = {
        "MC": {
            "name": "MC必选标准",
            "subcategories": {
                "mc1_product_leadership": {
                    "name": "产品团队领导力",
                    "description": "展示在产品开发或技术团队中的领导角色",
                    "keywords": ["CEO", "CTO", "创始人", "技术总监", "产品负责人", "团队领导", "主导", "带领团队", "研发负责人"]
                },
                "mc2_business_development": {
                    "name": "商业发展",
                    "description": "展示在商业拓展、融资、市场开发等方面的成就",
                    "keywords": ["融资", "投资", "商业合作", "市场拓展", "收入增长", "客户获取", "商业模式", "营收", "估值"]
                },
                "mc3_nonprofit": {
                    "name": "非营利组织",
                    "description": "在学术机构或非营利组织中的贡献",
                    "keywords": ["导师", "顾问", "志愿者", "公益", "学术指导", "孵化器", "创业导师", "评委"]
                },
                "mc4_expert_review": {
                    "name": "专家评审",
                    "description": "作为专家进行评审、审稿等活动",
                    "keywords": ["评审", "审稿", "评委", "专家", "顾问委员会", "评估", "咨询"]
                }
            }
        },
        "OC": {
            "name": "OC可选标准",
            "subcategories": {
                "oc1_innovation": {
                    "name": "创新",
                    "description": "技术创新、产品创新、专利等",
                    "keywords": ["创新", "专利", "发明", "突破", "首创", "独特", "原创", "技术突破", "新方法", "AI", "算法"]
                },
                "oc2_industry_recognition": {
                    "name": "行业认可",
                    "description": "获得的奖项、荣誉、媒体报道等",
                    "keywords": ["奖项", "荣誉", "认可", "表彰", "媒体报道", "行业奖", "获奖", "年度", "优秀"]
                },
                "oc3_significant_contribution": {
                    "name": "重大贡献",
                    "description": "对行业、公司或开源项目的重大贡献",
                    "keywords": ["重大贡献", "行业影响", "开源贡献", "标准制定", "技术推广", "产业影响", "里程碑"]
                },
                "oc4_academic": {
                    "name": "学术贡献",
                    "description": "论文发表、会议演讲、技术分享等",
                    "keywords": ["论文", "发表", "学术", "研究", "博士", "硕士", "会议", "演讲", "期刊", "引用"]
                }
            }
        },
        "RECOMMENDER": {
            "name": "推荐信息",
            "dynamic": True,  # 标记为动态类别，每个推荐人独立分组
            "subcategories": {
                # 动态生成：recommender_1, recommender_2, ... 每个推荐人一个独立分组
                # 格式: recommender_{序号}_{推荐人姓名}
            },
            "description": "每个推荐人的信息独立分组，包括：推荐人姓名、职位/头衔、所在机构、与申请人的关系、推荐理由、对申请人的评价等",
            "keywords": ["推荐人", "推荐信", "教授", "博士", "总监", "CEO", "专家", "背景", "关系", "合作", "共事", "指导", "认识", "了解", "评价", "推荐理由", "Professor", "Dr.", "Director"]
        }
    }
    
    def classify_content(self, project_id: str) -> Dict[str, Any]:
        """
        直接从原始材料提取并分类证据（使用 DAO 层）
        
        改进：不再从已提取的内容块分类，而是直接读取原始材料文件，
        这样可以保留完整的文件上下文，准确识别信息的主体人物（申请人vs推荐人）
        """
        try:
            # 获取项目的原始材料文件
            files = self._get_project_files(project_id)
            
            if not files:
                return {"success": False, "error": "没有找到项目材料文件"}
            
            total_files = len(files)
            
            # 记录分类进度
            self._update_classification_progress(project_id, {
                "status": "processing",
                "total_contents": total_files,
                "processed_contents": 0,
                "current_batch": 0,
                "total_batches": total_files,
                "total_classified": 0
            })
            
            self._log_extraction(
                project_id, 'classification', '开始从原始材料提取分类',
                status='started',
                response=f"共 {total_files} 个材料文件"
            )
            
            # 清除旧的分类结果
            self.classification_dao.delete_by_project(project_id)
            
            total_classified = 0
            classification_stats = {}
            
            # 记录每个文件的提取状态
            file_extraction_status = []
            
            # 逐个文件处理
            for file_idx, file_info in enumerate(files):
                filename = file_info.get('filename', '')
                file_category = file_info.get('category', '')
                file_type = file_info.get('file_type', '').lower()
                
                # 从扩展名推断类型
                if file_type in ['document', 'spreadsheet', 'image', 'other', '']:
                    ext = Path(filename).suffix.lower().lstrip('.')
                    if ext:
                        file_type = ext
                
                self._log_extraction(
                    project_id, 'classification', f'处理文件: {filename}',
                    response=f"文件 {file_idx + 1}/{total_files}"
                )
                
                # 读取文件原始内容
                file_content = self._read_file_for_classification(file_info)
                
                if not file_content:
                    # 记录提取失败的文件
                    reason = "扫描件（无文字层）" if file_type == 'pdf' else "无法读取文件内容"
                    file_extraction_status.append({
                        "filename": filename,
                        "status": "skipped",
                        "reason": reason,
                        "evidence_count": 0
                    })
                    logger.warning(f"无法读取文件内容: {filename}")
                    continue
                
                # 使用AI从原始材料直接提取并分类
                classifications = self._ai_extract_and_classify_from_file(
                    filename, file_content, file_category
                )
                
                # 记录提取成功的文件
                file_extraction_status.append({
                    "filename": filename,
                    "status": "extracted",
                    "reason": None,
                    "evidence_count": len(classifications)
                })
                
                # 保存分类结果
                for cls in classifications:
                    category = cls.get('category', '')
                    subcategory = cls.get('subcategory', '')
                    content = cls.get('content', '')
                    subject_person = cls.get('subject', 'applicant')
                    relevance = cls.get('relevance_score', 0.8)
                    evidence_type = cls.get('evidence_type', '')
                    key_points = json.dumps(cls.get('key_points', []), ensure_ascii=False)
                    
                    # 推荐人相关字段
                    recommender_name = cls.get('recommender_name', '')
                    recommender_title = cls.get('recommender_title', '')
                    recommender_org = cls.get('recommender_org', '')
                    relationship = cls.get('relationship', '')
                    
                    if content and category and subcategory:
                        self.classification_dao.add_classification(
                            project_id=project_id,
                            category=category,
                            subcategory=subcategory,
                            content=content,
                            source_file=filename,
                            source_page=cls.get('source_page'),
                            relevance_score=relevance,
                            evidence_type=evidence_type,
                            key_points=key_points,
                            subject_person=subject_person,
                            recommender_name=recommender_name,
                            recommender_title=recommender_title,
                            recommender_org=recommender_org,
                            relationship=relationship
                        )
                        
                        total_classified += 1
                        
                        # 统计（推荐人按姓名分组统计）
                        if category == 'RECOMMENDER' and recommender_name:
                            key = f"RECOMMENDER_{recommender_name}"
                        else:
                            key = f"{category}_{subcategory}"
                        classification_stats[key] = classification_stats.get(key, 0) + 1
                
                # 更新进度
                self._update_classification_progress(project_id, {
                    "status": "processing",
                    "total_contents": total_files,
                    "processed_contents": file_idx + 1,
                    "current_batch": file_idx + 1,
                    "total_batches": total_files,
                    "total_classified": total_classified
                })
                
                logger.info(f"分类进度: 文件 {file_idx + 1}/{total_files} ({filename}), 已分类 {total_classified} 条")
            
            # 统计文件提取情况
            extracted_files = [f for f in file_extraction_status if f['status'] == 'extracted']
            skipped_files = [f for f in file_extraction_status if f['status'] == 'skipped']
            
            # 更新最终状态（包含文件提取情况）
            self._update_classification_progress(project_id, {
                "status": "completed",
                "total_contents": total_files,
                "processed_contents": total_files,
                "current_batch": total_files,
                "total_batches": total_files,
                "total_classified": total_classified
            })
            
            self._log_extraction(
                project_id, 'classification', '内容分类完成',
                response=json.dumps({
                    "total_files": total_files,
                    "extracted_files": len(extracted_files),
                    "skipped_files": len(skipped_files),
                    "total_classified": total_classified,
                    "stats": classification_stats
                }, ensure_ascii=False)
            )
            
            return {
                "success": True,
                "data": {
                    "total_files": total_files,
                    "extracted_files": len(extracted_files),
                    "skipped_files": len(skipped_files),
                    "total_classified": total_classified,
                    "stats": classification_stats,
                    "file_status": file_extraction_status
                }
            }
            
        except Exception as e:
            logger.error(f"内容分类失败: {e}")
            import traceback
            traceback.print_exc()
            # 更新失败状态
            try:
                self._update_classification_progress(project_id, {
                    "status": "failed",
                    "error": str(e)
                })
            except:
                pass
            self._log_extraction(
                project_id, 'classification', '内容分类失败',
                status='error', error_message=str(e)
            )
            return {"success": False, "error": str(e)}
    
    def _read_file_for_classification(self, file_info: Dict) -> str:
        """
        读取文件内容用于分类
        
        使用统一文件存储接口，自动支持任何存储后端
        """
        file_type = file_info.get('file_type', '').lower()
        filename = file_info.get('filename', '')
        
        # 从扩展名推断类型
        if file_type in ['document', 'spreadsheet', 'image', 'other', '']:
            ext = Path(filename).suffix.lower().lstrip('.')
            if ext:
                file_type = ext
        
        try:
            # 使用统一接口获取文件内容（自动处理任何存储类型）
            file_bytes = self._get_file_bytes(file_info, filename)
            if file_bytes is None:
                return ""
            return self._read_content_from_bytes(file_bytes, file_type, filename)
                
        except Exception as e:
            logger.error(f"读取文件失败 {filename}: {e}")
            return ""
    
    def _read_content_from_bytes(self, file_bytes: bytes, file_type: str, filename: str) -> str:
        """从内存字节流读取文件内容（用于分类）"""
        from io import BytesIO
        
        content_parts = []
        
        try:
            if file_type == 'pdf':
                import fitz
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                total_pages = len(doc)
                pages_with_text = 0
                for page_num, page in enumerate(doc):
                    text = page.get_text()
                    if text.strip():
                        content_parts.append(f"[第{page_num+1}页]\n{text}")
                        pages_with_text += 1
                doc.close()
                
                # 如果是扫描件（有页面但没有文字），记录日志
                if total_pages > 0 and pages_with_text == 0:
                    logger.warning(f"PDF 可能是扫描件（无文字层）: {filename}, 共 {total_pages} 页")
                else:
                    logger.info(f"PDF 解析完成: {filename}, {pages_with_text}/{total_pages} 页有文字")
                
            elif file_type in ['docx', 'doc']:
                import docx
                doc = docx.Document(BytesIO(file_bytes))
                for para in doc.paragraphs:
                    if para.text.strip():
                        content_parts.append(para.text)
                # 也读取表格内容
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                        if row_text.strip():
                            content_parts.append(row_text)
                        
            elif file_type in ['xlsx', 'xls']:
                import pandas as pd
                xls = pd.ExcelFile(BytesIO(file_bytes))
                for sheet_name in xls.sheet_names:
                    df = pd.read_excel(xls, sheet_name=sheet_name)
                    content_parts.append(f"[工作表: {sheet_name}]\n{df.to_string()}")
                    
            elif file_type in ['txt', 'md', 'json']:
                # 尝试多种编码
                content = None
                for encoding in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
                    try:
                        content = file_bytes.decode(encoding)
                        break
                    except UnicodeDecodeError:
                        continue
                if content:
                    content_parts.append(content)
            
            return '\n\n'.join(content_parts)
            
        except Exception as e:
            logger.error(f"从内存读取文件内容失败 {filename}: {e}")
            return ""
    
    def _read_content_from_file(self, file_path: str, file_type: str, filename: str) -> str:
        """从本地文件读取内容（用于分类）"""
        content_parts = []
        
        try:
            if file_type == 'pdf':
                import fitz
                doc = fitz.open(file_path)
                for page_num, page in enumerate(doc):
                    text = page.get_text()
                    if text.strip():
                        content_parts.append(f"[第{page_num+1}页]\n{text}")
                doc.close()
                
            elif file_type in ['docx', 'doc']:
                import docx
                doc = docx.Document(file_path)
                for para in doc.paragraphs:
                    if para.text.strip():
                        content_parts.append(para.text)
                        
            elif file_type in ['xlsx', 'xls']:
                import pandas as pd
                xls = pd.ExcelFile(file_path)
                for sheet_name in xls.sheet_names:
                    df = pd.read_excel(xls, sheet_name=sheet_name)
                    content_parts.append(f"[工作表: {sheet_name}]\n{df.to_string()}")
                    
            elif file_type in ['txt', 'md', 'json']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content_parts.append(f.read())
            
            return '\n\n'.join(content_parts)
            
        except Exception as e:
            logger.error(f"读取本地文件内容失败 {filename}: {e}")
            return ""
    
    def _ai_extract_and_classify_from_file(self, filename: str, file_content: str, file_category: str) -> List[Dict]:
        """
        使用AI从原始材料直接提取并分类证据
        
        关键改进：
        1. 在提取时识别信息的主体人物（申请人/推荐人）
        2. 每个推荐人独立分组，使用推荐人姓名作为标识
        """
        if not self.llm_client or not file_content:
            return []
        
        try:
            # 构建分类说明（不包含动态的RECOMMENDER类别）
            categories_desc = []
            for cat_key, cat_info in self.CLASSIFICATION_CATEGORIES.items():
                if cat_key == 'RECOMMENDER':
                    continue  # 推荐人类别单独处理
                for sub_key, sub_info in cat_info['subcategories'].items():
                    categories_desc.append(f"- {cat_key}/{sub_key}: {sub_info['name']} - {sub_info['description']}")
            
            # 根据文件名和类别提供上下文提示
            file_context = ""
            is_recommender_file = '推荐' in filename or 'recommender' in filename.lower() or 'rl' in filename.lower()
            if is_recommender_file:
                file_context = "【注意】这是推荐人相关文件，其中描述的个人经历和背景是推荐人的，不是申请人的。请识别出推荐人的姓名。"
            elif 'CV' in filename or '简历' in filename:
                file_context = "【注意】这是申请人的简历，其中的经历和成就属于申请人本人。"
            
            # 限制内容长度
            content_preview = file_content[:10000] if len(file_content) > 10000 else file_content
            
            prompt = f"""请从以下材料中提取GTV签证申请的相关证据，并进行分类。

## 文件信息
- 文件名: {filename}
- 文件类别: {file_category or '未知'}
{file_context}

## 重要：识别信息主体
每条证据必须准确识别其描述的是谁：
- "applicant": 申请人本人的经历、成就、工作、教育等
- "recommender": 推荐人的背景、职位、学历、成就、与申请人的关系等
- "other": 公司、组织或其他人物的信息

⚠️ 特别注意推荐人信息处理：
- 每个推荐人必须独立识别，记录其姓名
- 不同推荐人的信息绝对不能混淆
- 从材料中提取有多少个推荐人就记录多少个，不要限制数量
- 推荐人信息包括：姓名、职位、所在机构、与申请人的关系、对申请人的评价、推荐理由等

## 可选分类（申请人相关）
{chr(10).join(categories_desc)}

## 推荐人分类（每个推荐人独立分组）
对于推荐人相关信息，使用以下格式：
- category: "RECOMMENDER"
- subcategory: "recommender_{{推荐人姓名}}" （例如：recommender_张三、recommender_Dr_John_Smith）
- recommender_name: 推荐人的完整姓名（必填，用于分组）

## 材料内容
{content_preview}

## 输出要求
提取所有与GTV申请相关的证据，返回JSON数组。

对于申请人相关证据：
- category: MC/OC
- subcategory: 子类别代码
- subject: "applicant"
- content: 证据内容（100-500字）
- relevance_score: 相关度（0.6-1.0）
- evidence_type: 证据类型
- key_points: 2-3个关键词

对于推荐人相关证据（每个推荐人独立记录）：
- category: "RECOMMENDER"
- subcategory: "recommender_{{推荐人姓名}}"
- subject: "recommender"
- recommender_name: 推荐人完整姓名（必填）
- recommender_title: 推荐人职位/头衔
- recommender_org: 推荐人所在机构
- relationship: 与申请人的关系
- content: 推荐人信息或评价内容（100-500字）
- relevance_score: 相关度（0.6-1.0）
- evidence_type: 证据类型（如"推荐人背景"、"推荐评价"、"合作关系"等）
- key_points: 2-3个关键词

如果材料中没有相关证据，返回空数组[]。

## 输出格式（只返回JSON数组）
```json
[
  {{"category": "MC", "subcategory": "mc1_product_leadership", "subject": "applicant", "content": "申请人作为创始人兼CEO...", "relevance_score": 0.9, "evidence_type": "工作经历", "key_points": ["创始人", "CEO"]}},
  {{"category": "RECOMMENDER", "subcategory": "recommender_张三", "subject": "recommender", "recommender_name": "张三", "recommender_title": "教授", "recommender_org": "北京大学", "relationship": "博士导师", "content": "张三教授是北京大学计算机系教授，曾指导申请人完成博士论文...", "relevance_score": 0.9, "evidence_type": "推荐人背景", "key_points": ["教授", "北京大学", "博士导师"]}},
  {{"category": "RECOMMENDER", "subcategory": "recommender_Dr_John_Smith", "subject": "recommender", "recommender_name": "Dr. John Smith", "recommender_title": "CEO", "recommender_org": "Tech Corp", "relationship": "前雇主", "content": "Dr. John Smith是Tech Corp的CEO，申请人在其公司工作期间...", "relevance_score": 0.85, "evidence_type": "推荐评价", "key_points": ["CEO", "Tech Corp", "前雇主"]}}
]
```"""

            response = self.llm_client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.2,
                max_tokens=4000
            )
            
            result_text = response.choices[0].message.content.strip()
            
            # 提取JSON
            if '```json' in result_text:
                result_text = result_text.split('```json')[1].split('```')[0].strip()
            elif '```' in result_text:
                parts = result_text.split('```')
                if len(parts) >= 2:
                    result_text = parts[1].strip()
                    if result_text.startswith('json'):
                        result_text = result_text[4:].strip()
            
            # 找到JSON数组
            start_idx = result_text.find('[')
            end_idx = result_text.rfind(']')
            if start_idx != -1 and end_idx != -1:
                result_text = result_text[start_idx:end_idx+1]
            
            classifications = json.loads(result_text)
            
            # 验证结果
            valid_classifications = []
            for cls in classifications:
                if isinstance(cls, dict) and cls.get('content') and cls.get('category'):
                    # 确保必要字段
                    cls['subject'] = cls.get('subject', 'applicant')
                    cls['relevance_score'] = cls.get('relevance_score', 0.7)
                    valid_classifications.append(cls)
            
            logger.info(f"从文件 {filename} 提取了 {len(valid_classifications)} 条证据")
            return valid_classifications
            
        except Exception as e:
            logger.error(f"AI提取分类失败 {filename}: {e}")
            return []
    
    def _ai_batch_classify(self, batch_items: List[Dict]) -> Dict[str, List[Dict]]:
        """
        使用AI批量分类多个内容块（一次LLM调用处理多个内容）
        
        Args:
            batch_items: 内容块列表，每个包含 id, source_file, source_page, content
            
        Returns:
            按内容ID索引的分类结果字典
        """
        results = {}
        
        if not self.llm_client:
            # 无LLM时使用关键词匹配
            for item in batch_items:
                results[str(item['id'])] = self._keyword_classify_content(item['content'], item['source_file'])
            return results
        
        try:
            # 构建分类说明
            categories_desc = []
            for cat_key, cat_info in self.CLASSIFICATION_CATEGORIES.items():
                for sub_key, sub_info in cat_info['subcategories'].items():
                    categories_desc.append(f"- {cat_key}/{sub_key}: {sub_info['name']}")
            
            # 构建批量内容
            contents_text = []
            for i, item in enumerate(batch_items):
                contents_text.append(f"""[内容{item['id']}] 来源: {item['source_file']}
{item['content'][:600]}
---""")
            
            prompt = f"""请分析以下多个内容块，判断每个属于GTV签证申请的哪个证据类别。

## 重要：识别信息主体
每条信息必须识别其描述的主体人物：
- "applicant": 申请人本人的经历、成就、背景
- "recommender": 推荐人的背景、资质、职位、与申请人的关系等
- "other": 其他人物或组织的信息

注意区分：
- 申请人的CV、工作经历、成就属于 applicant
- 推荐人的介绍、推荐人的学历职位属于 recommender
- 推荐人信息文件中描述的推荐人背景属于 recommender，不要与申请人混淆

⚠️ 推荐人信息特别处理：
- 每个推荐人必须独立识别，使用推荐人姓名作为标识
- 不同推荐人的信息不能混淆
- subcategory格式为: recommender_{{推荐人姓名}}

## 可选类别（申请人相关）
{chr(10).join(categories_desc)}

## 推荐人分类
对于推荐人，subcategory使用格式: recommender_{{推荐人姓名}}
例如: recommender_张三, recommender_Dr_John_Smith

## 待分类内容
{chr(10).join(contents_text)}

## 输出要求
返回一个JSON对象，key是内容ID，value是该内容的分类数组。

对于申请人相关分类:
- category: MC/OC
- subcategory: 子类别代码
- subject: "applicant"
- relevance_score: 相关度0.5-1.0
- evidence_type: 证据类型
- key_points: 2-3个关键词

对于推荐人相关分类:
- category: "RECOMMENDER"
- subcategory: "recommender_{{推荐人姓名}}"
- subject: "recommender"
- recommender_name: 推荐人完整姓名
- relevance_score: 相关度0.5-1.0
- evidence_type: 证据类型
- key_points: 2-3个关键词

如果某内容不属于任何类别，对应的value为空数组[]。

## 输出格式（严格按此格式，只返回JSON）
```json
{{
  "123": [{{"category": "MC", "subcategory": "mc1_product_leadership", "subject": "applicant", "relevance_score": 0.9, "evidence_type": "工作经历", "key_points": ["创始人", "CEO"]}}],
  "124": [],
  "125": [{{"category": "RECOMMENDER", "subcategory": "recommender_张三", "subject": "recommender", "recommender_name": "张三", "relevance_score": 0.8, "evidence_type": "推荐人背景", "key_points": ["教授", "北京大学"]}}]
}}
```"""

            response = self.llm_client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.2,
                max_tokens=2000
            )
            
            result_text = response.choices[0].message.content.strip()
            
            # 提取JSON
            if '```json' in result_text:
                result_text = result_text.split('```json')[1].split('```')[0].strip()
            elif '```' in result_text:
                parts = result_text.split('```')
                if len(parts) >= 2:
                    result_text = parts[1].strip()
                    # 移除可能的语言标识符
                    if result_text.startswith('json'):
                        result_text = result_text[4:].strip()
            
            # 尝试找到JSON对象
            start_idx = result_text.find('{')
            end_idx = result_text.rfind('}')
            if start_idx != -1 and end_idx != -1:
                result_text = result_text[start_idx:end_idx+1]
            
            batch_results = json.loads(result_text)
            
            # 验证和过滤结果
            for content_id, classifications in batch_results.items():
                valid_classifications = []
                if isinstance(classifications, list):
                    for cls in classifications:
                        if isinstance(cls, dict) and cls.get('relevance_score', 0) >= 0.5:
                            valid_classifications.append(cls)
                results[str(content_id)] = valid_classifications
            
            # 确保所有输入内容都有结果
            for item in batch_items:
                if str(item['id']) not in results:
                    results[str(item['id'])] = []
            
            return results
            
        except Exception as e:
            logger.warning(f"AI批量分类失败，回退到关键词匹配: {e}")
            # 回退到关键词匹配
            for item in batch_items:
                results[str(item['id'])] = self._keyword_classify_content(item['content'], item['source_file'])
            return results
    
    def _keyword_classify_content(self, content: str, source_file: str) -> List[Dict]:
        """
        使用关键词匹配进行分类（备用方案）
        """
        classifications = []
        content_lower = content.lower()
        
        for cat_key, cat_info in self.CLASSIFICATION_CATEGORIES.items():
            for sub_key, sub_info in cat_info['subcategories'].items():
                keywords = sub_info.get('keywords', [])
                matched_keywords = [kw for kw in keywords if kw.lower() in content_lower]
                
                if len(matched_keywords) >= 2:  # 至少匹配2个关键词
                    relevance = min(len(matched_keywords) / len(keywords), 1.0)
                    classifications.append({
                        "category": cat_key,
                        "subcategory": sub_key,
                        "relevance_score": round(relevance * 0.8, 2),  # 关键词匹配最高0.8
                        "evidence_type": "关键词匹配",
                        "key_points": matched_keywords[:3]
                    })
        
        return classifications
    
    def get_classifications(self, project_id: str, category: str = None) -> Dict[str, Any]:
        """获取项目的分类结果（使用 DAO 层）"""
        try:
            # 使用 DAO 获取分类
            if category:
                rows = self.classification_dao.get_with_details(project_id, category)
            else:
                rows = self.classification_dao.get_all_by_project(project_id)
            
            # 按类别和子类别组织结果
            result = {}
            for row in rows:
                cat = row.get('category')
                subcat = row.get('subcategory')
                
                if cat not in result:
                    cat_info = self.CLASSIFICATION_CATEGORIES.get(cat, {})
                    result[cat] = {
                        "name": cat_info.get('name', cat),
                        "dynamic": cat_info.get('dynamic', False),
                        "subcategories": {}
                    }
                
                if subcat not in result[cat]['subcategories']:
                    if cat == 'RECOMMENDER':
                        recommender_name = row.get('recommender_name') or subcat.replace('recommender_', '').replace('_', ' ')
                        result[cat]['subcategories'][subcat] = {
                            "name": f"推荐人：{recommender_name}",
                            "description": f"{row.get('recommender_title') or ''} - {row.get('recommender_org') or ''}".strip(' - '),
                            "recommender_name": recommender_name,
                            "recommender_title": row.get('recommender_title') or '',
                            "recommender_org": row.get('recommender_org') or '',
                            "items": []
                        }
                    else:
                        subcat_info = self.CLASSIFICATION_CATEGORIES.get(cat, {}).get('subcategories', {}).get(subcat, {})
                        result[cat]['subcategories'][subcat] = {
                            "name": subcat_info.get('name', subcat),
                            "description": subcat_info.get('description', ''),
                            "items": []
                        }
                
                # 解析 key_points
                key_points = []
                try:
                    kp = row.get('key_points')
                    key_points = json.loads(kp) if kp else []
                except:
                    pass
                
                item_data = {
                    "id": row['id'],
                    "content": row.get('content'),
                    "source_file": row.get('source_file'),
                    "source_page": row.get('source_page'),
                    "relevance_score": row.get('relevance_score'),
                    "evidence_type": row.get('evidence_type'),
                    "key_points": key_points,
                    "subject_person": row.get('subject_person') or 'applicant'
                }
                
                if cat == 'RECOMMENDER':
                    item_data["recommender_name"] = row.get('recommender_name') or ''
                    item_data["recommender_title"] = row.get('recommender_title') or ''
                    item_data["recommender_org"] = row.get('recommender_org') or ''
                    item_data["relationship"] = row.get('relationship') or ''
                
                result[cat]['subcategories'][subcat]['items'].append(item_data)
            
            return {
                "success": True,
                "data": {
                    "classifications": result,
                    "total_items": len(rows)
                }
            }
            
        except Exception as e:
            logger.error(f"获取分类结果失败: {e}")
            return {"success": False, "error": str(e)}
    
    def get_classification_summary(self, project_id: str) -> Dict[str, Any]:
        """获取分类统计摘要（使用 DAO 层）"""
        try:
            rows = self.classification_dao.get_classification_summary(project_id)
            
            summary = {}
            recommender_list = []
            
            for row in rows:
                cat = row['category']
                subcat = row['subcategory']
                
                if cat not in summary:
                    cat_info = self.CLASSIFICATION_CATEGORIES.get(cat, {})
                    summary[cat] = {
                        "name": cat_info.get('name', cat),
                        "total": 0,
                        "subcategories": {},
                        "dynamic": cat_info.get('dynamic', False)
                    }
                
                if cat == 'RECOMMENDER':
                    recommender_name = row.get('recommender_name') or subcat.replace('recommender_', '').replace('_', ' ')
                    subcat_display_name = f"推荐人：{recommender_name}"
                    
                    summary[cat]['subcategories'][subcat] = {
                        "name": subcat_display_name,
                        "recommender_name": recommender_name,
                        "recommender_title": row.get('recommender_title') or '',
                        "recommender_org": row.get('recommender_org') or '',
                        "count": row['count'],
                        "avg_score": round(row['avg_score'], 2) if row.get('avg_score') else 0
                    }
                    
                    recommender_list.append({
                        "name": recommender_name,
                        "title": row.get('recommender_title') or '',
                        "org": row.get('recommender_org') or '',
                        "evidence_count": row['count']
                    })
                else:
                    subcat_info = self.CLASSIFICATION_CATEGORIES.get(cat, {}).get('subcategories', {}).get(subcat, {})
                    summary[cat]['subcategories'][subcat] = {
                        "name": subcat_info.get('name', subcat),
                        "count": row['count'],
                        "avg_score": round(row['avg_score'], 2) if row.get('avg_score') else 0
                    }
                
                summary[cat]['total'] += row['count']
            
            return {
                "success": True,
                "data": {
                    "summary": summary,
                    "recommenders": recommender_list,
                    "recommender_count": len(recommender_list)
                }
            }
            
        except Exception as e:
            logger.error(f"获取分类摘要失败: {e}")
            return {"success": False, "error": str(e)}

    def _update_classification_progress(self, project_id: str, progress_data: Dict):
        """更新分类进度到数据库（使用 DAO 层）"""
        try:
            self.classification_dao.update_progress_data(project_id, progress_data)
        except Exception as e:
            logger.warning(f"更新分类进度失败: {e}")

    def get_classification_progress(self, project_id: str) -> Dict[str, Any]:
        """获取分类进度（使用 DAO 层）"""
        try:
            row = self.classification_dao.get_progress(project_id)
            
            if row:
                progress_pct = 0
                total = row.get('total_contents', 0) or 0
                processed = row.get('processed_contents', 0) or 0
                if total > 0:
                    progress_pct = round(processed / total * 100, 1)
                
                return {
                    "success": True,
                    "data": {
                        "status": row.get('status'),
                        "total_contents": total,
                        "processed_contents": processed,
                        "current_batch": row.get('current_batch', 0),
                        "total_batches": row.get('total_batches', 0),
                        "total_classified": row.get('total_classified', 0),
                        "progress_percent": progress_pct,
                        "error": row.get('error'),
                        "updated_at": row.get('updated_at')
                    }
                }
            else:
                return {
                    "success": True,
                    "data": {
                        "status": "idle",
                        "total_contents": 0,
                        "processed_contents": 0,
                        "progress_percent": 0
                    }
                }
            
        except Exception as e:
            logger.error(f"获取分类进度失败: {e}")
            return {"success": False, "error": str(e)}

    def update_classification(self, project_id: str, classification_id: int, data: Dict) -> Dict[str, Any]:
        """更新单个分类项（使用 DAO 层）"""
        try:
            # 处理 key_points 字段
            key_points = data.get('key_points')
            if isinstance(key_points, list):
                key_points = json.dumps(key_points, ensure_ascii=False)
            
            success = self.classification_dao.update_classification(
                classification_id,
                category=data.get('category'),
                subcategory=data.get('subcategory'),
                content=data.get('content'),
                confidence=data.get('relevance_score')
            )
            
            if not success:
                return {"success": False, "error": "分类项不存在或无需更新"}
            
            return {"success": True, "message": "更新成功"}
            
        except Exception as e:
            logger.error(f"更新分类失败: {e}")
            return {"success": False, "error": str(e)}

    def delete_classification(self, project_id: str, classification_id: int) -> Dict[str, Any]:
        """删除单个分类项（使用 DAO 层）"""
        try:
            success = self.classification_dao.delete_by_id(classification_id)
            
            if not success:
                return {"success": False, "error": "分类项不存在"}
            
            return {"success": True, "message": "删除成功"}
            
        except Exception as e:
            logger.error(f"删除分类失败: {e}")
            return {"success": False, "error": str(e)}

    def add_classification(self, project_id: str, data: Dict) -> Dict[str, Any]:
        """手动添加分类项（使用 DAO 层）"""
        try:
            # 必填字段检查
            category = data.get('category')
            subcategory = data.get('subcategory')
            content = data.get('content')
            
            if not all([category, subcategory, content]):
                return {"success": False, "error": "缺少必填字段: category, subcategory, content"}
            
            # 验证类别有效性
            if category not in self.CLASSIFICATION_CATEGORIES:
                return {"success": False, "error": f"无效的类别: {category}"}
            
            # 推荐人类别是动态的，允许 recommender_xxx 格式的子类别
            if category == 'RECOMMENDER':
                if not subcategory.startswith('recommender_'):
                    return {"success": False, "error": f"推荐人子类别格式应为: recommender_{{推荐人姓名}}"}
            elif subcategory not in self.CLASSIFICATION_CATEGORIES[category]['subcategories']:
                return {"success": False, "error": f"无效的子类别: {subcategory}"}
            
            key_points = data.get('key_points', [])
            if isinstance(key_points, list):
                key_points = json.dumps(key_points, ensure_ascii=False)
            
            new_id = self.classification_dao.add_classification(
                project_id=project_id,
                category=category,
                subcategory=subcategory,
                content=content,
                source_file=data.get('source_file', '手动添加'),
                source_page=data.get('source_page'),
                relevance_score=data.get('relevance_score', 1.0),
                evidence_type=data.get('evidence_type', '用户添加'),
                key_points=key_points,
                subject_person=data.get('subject_person', 'applicant'),
                recommender_name=data.get('recommender_name', ''),
                recommender_title=data.get('recommender_title', ''),
                recommender_org=data.get('recommender_org', ''),
                relationship=data.get('relationship', '')
            )
            
            return {
                "success": True,
                "data": {"id": new_id},
                "message": "添加成功"
            }
            
        except Exception as e:
            logger.error(f"添加分类失败: {e}")
            return {"success": False, "error": str(e)}

    def get_classified_evidence_for_framework(self, project_id: str) -> Dict[str, Any]:
        """
        获取整理好的分类证据，用于构建GTV申请框架（使用 DAO 层）
        
        返回按MC/OC标准组织的证据数据，包含原始内容和来源信息
        """
        try:
            rows = self.classification_dao.get_classified_evidence(project_id)
            
            if not rows:
                return {"success": False, "error": "没有分类证据，请先进行内容分类"}
            
            # 按标准组织证据
            evidence = {
                "MC": {},
                "OC": {},
                "RECOMMENDER": {}
            }
            
            for row in rows:
                cat = row.get('category', '')
                subcat = row.get('subcategory', '')
                
                if cat not in evidence:
                    evidence[cat] = {}
                
                if subcat not in evidence[cat]:
                    subcat_info = self.CLASSIFICATION_CATEGORIES.get(cat, {}).get('subcategories', {}).get(subcat, {})
                    evidence[cat][subcat] = {
                        "name": subcat_info.get('name', subcat),
                        "items": []
                    }
                
                key_points = []
                try:
                    key_points = json.loads(row.get('key_points', '[]')) if row.get('key_points') else []
                except:
                    pass
                
                evidence[cat][subcat]["items"].append({
                    "id": row.get('id'),
                    "content": row.get('content', ''),
                    "source_file": row.get('source_file', ''),
                    "source_page": row.get('source_page'),
                    "relevance_score": row.get('relevance_score', 0),
                    "evidence_type": row.get('evidence_type', ''),
                    "key_points": key_points
                })
            
            return {
                "success": True,
                "data": {
                    "evidence": evidence,
                    "total_items": len(rows)
                }
            }
            
        except Exception as e:
            logger.error(f"获取分类证据失败: {e}")
            return {"success": False, "error": str(e)}


# 测试
if __name__ == "__main__":
    agent = ContentExtractionAgent("copywriting.db", "./uploads")
    
    # 测试提取
    result = agent.extract_project_files("TEST001")
    print("提取结果:", json.dumps(result, ensure_ascii=False, indent=2))
    
    # 测试获取上下文
    context = agent.get_project_context("TEST001")
    print("\n上下文:", context.get("data", {}).get("context", "")[:500])
