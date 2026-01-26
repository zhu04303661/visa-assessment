#!/usr/bin/env python3
"""
GTV签证申请原始材料收集管理系统
基于材料要求清单设计，帮助客户系统性地收集和管理申请所需的原始材料
"""

import sqlite3
import json
import os
import uuid
from datetime import datetime
from typing import Dict, Any, List, Optional
from contextlib import contextmanager
from pathlib import Path

from logger_config import setup_module_logger

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = setup_module_logger("raw_material_manager", os.getenv("LOG_LEVEL", "INFO"))


# ==================== 材料分类定义 ====================
# 基于GTV申请材料要求清单

MATERIAL_CATEGORIES = {
    "folder_1": {
        "name": "申请人个人资料",
        "name_en": "Personal Documents",
        "description": "申请人的基本身份和资质证明材料",
        "order": 1,
        "items": [
            {
                "item_id": "resume",
                "name": "个人简历",
                "name_en": "Resume/CV",
                "description": "详细的工作经历和教育背景",
                "required": True,
                "file_types": ["pdf", "docx", "doc"],
                "has_form": False,
                "tips": "建议使用专业格式，突出技术成就和领导经验"
            },
            {
                "item_id": "passport",
                "name": "护照",
                "name_en": "Passport",
                "description": "护照个人信息页、签证页（如有）",
                "required": True,
                "file_types": ["pdf", "jpg", "png"],
                "has_form": False,
                "tips": "需清晰扫描件，包括个人信息页和所有签证页"
            },
            {
                "item_id": "education_cert",
                "name": "学历学位证书",
                "name_en": "Education Certificates",
                "description": "本科及以上学历、学位证书",
                "required": True,
                "file_types": ["pdf", "jpg", "png"],
                "has_form": False,
                "multiple": True,
                "tips": "需提供所有高等教育学历证书，如有海外学历需提供认证"
            },
            {
                "item_id": "professional_cert",
                "name": "专业资格证书",
                "name_en": "Professional Certificates",
                "description": "行业认证、专业资格等证书",
                "required": False,
                "file_types": ["pdf", "jpg", "png"],
                "has_form": False,
                "multiple": True,
                "tips": "如有行业认证（AWS、Google Cloud等）请一并提供"
            },
            {
                "item_id": "personal_statement",
                "name": "个人陈述",
                "name_en": "Personal Statement",
                "description": "由顾问协助撰写的个人陈述",
                "required": True,
                "file_types": ["pdf", "docx"],
                "has_form": False,
                "generated": True,
                "tips": "这将由顾问根据您的材料协助撰写"
            },
            {
                "item_id": "other_docs",
                "name": "其他文档",
                "name_en": "Other Documents",
                "description": "未分类的其他相关文档，可后续手动调整分类",
                "required": False,
                "file_types": ["pdf", "docx", "doc", "xlsx", "xls", "pptx", "ppt", "jpg", "jpeg", "png", "txt"],
                "has_form": False,
                "multiple": True,
                "tips": "系统未能自动识别的文件会暂存于此，请在文件列表中手动调整分类"
            }
        ]
    },
    "folder_2": {
        "name": "当前雇主材料",
        "name_en": "Current Employer Documents",
        "description": "当前工作单位的相关证明材料",
        "order": 2,
        "items": [
            {
                "item_id": "employment_info_form",
                "name": "就职信息采集表",
                "name_en": "Employment Information Form",
                "description": "详细的工作职责、项目经历等信息",
                "required": True,
                "file_types": ["docx", "pdf"],
                "has_form": True,
                "form_type": "employment_info",
                "tips": "请尽量详细填写，特别是主导的项目和取得的成果"
            },
            {
                "item_id": "employment_letter",
                "name": "雇佣证明信",
                "name_en": "Employment Verification Letter",
                "description": "公司出具的正式雇佣证明",
                "required": True,
                "file_types": ["pdf"],
                "has_form": False,
                "generated": True,
                "tips": "需使用公司抬头信纸，由HR部门签字盖章"
            },
            {
                "item_id": "supervisor_letter",
                "name": "领导推荐信",
                "name_en": "Supervisor Reference Letter",
                "description": "直属领导或高管出具的推荐信",
                "required": True,
                "file_types": ["pdf"],
                "has_form": False,
                "generated": True,
                "tips": "由顾问协助草拟，需领导签字"
            },
            {
                "item_id": "income_proof",
                "name": "收入证明",
                "name_en": "Income Certificate",
                "description": "工资单、纳税证明等收入相关材料",
                "required": True,
                "file_types": ["pdf", "jpg", "png"],
                "has_form": False,
                "multiple": True,
                "tips": "可提供工资单、银行流水或纳税证明"
            },
            {
                "item_id": "org_chart",
                "name": "组织架构图",
                "name_en": "Organization Chart",
                "description": "显示您在公司中职位的组织架构图",
                "required": False,
                "file_types": ["pdf", "jpg", "png"],
                "has_form": False,
                "tips": "标注您的位置，展示管理范围"
            },
            {
                "item_id": "company_awards",
                "name": "公司获奖/表彰",
                "name_en": "Company Awards",
                "description": "公司获得的行业奖项、荣誉证书等",
                "required": False,
                "file_types": ["pdf", "jpg", "png"],
                "has_form": False,
                "multiple": True,
                "tips": "如：国家级知识产权示范企业、行业排名等"
            },
            {
                "item_id": "personal_awards",
                "name": "个人获奖/表彰",
                "name_en": "Personal Awards",
                "description": "您在公司获得的个人奖项、表彰邮件等",
                "required": False,
                "file_types": ["pdf", "jpg", "png"],
                "has_form": False,
                "multiple": True,
                "tips": "如：优秀员工、技术贡献奖等"
            },
            {
                "item_id": "company_brochure",
                "name": "公司宣传材料",
                "name_en": "Company Brochure",
                "description": "公司宣传册、产品介绍等",
                "required": False,
                "file_types": ["pdf", "jpg", "png"],
                "has_form": False,
                "multiple": True,
                "tips": "纸质宣传册请提供四边完整的扫描件"
            },
            {
                "item_id": "company_media",
                "name": "公司媒体报道",
                "name_en": "Company Media Coverage",
                "description": "知名媒体对公司的报道链接",
                "required": False,
                "file_types": [],
                "has_form": True,
                "form_type": "media_links",
                "tips": "提供媒体报道的网址链接"
            }
        ]
    },
    "folder_3": {
        "name": "过往雇主材料",
        "name_en": "Previous Employer Documents",
        "description": "过往工作单位的材料（可添加多家公司）",
        "order": 3,
        "is_repeatable": True,  # 整个分类可重复添加
        "repeat_label": "公司",
        "min_count": 1,
        "max_count": 10,
        "items": [
            {
                "item_id": "prev_company_name",
                "name": "公司名称",
                "name_en": "Company Name",
                "description": "过往雇主公司名称",
                "required": True,
                "file_types": [],
                "has_form": True,
                "form_type": "prev_company_basic",
                "tips": "填写公司名称、职位、时间段"
            },
            {
                "item_id": "prev_employment_form",
                "name": "就职信息采集表",
                "name_en": "Employment Info Form",
                "description": "详细的工作职责和项目经历",
                "required": True,
                "file_types": ["docx", "pdf"],
                "has_form": True,
                "form_type": "prev_employment",
                "tips": "详细描述工作职责和主要成就"
            },
            {
                "item_id": "prev_employment_proof",
                "name": "工作证明/在职证明",
                "name_en": "Employment Certificate",
                "description": "公司出具的工作证明或在职证明",
                "required": False,
                "file_types": ["pdf", "jpg", "png"],
                "has_form": False,
                "tips": "如有正式的工作证明请上传"
            },
            {
                "item_id": "prev_awards",
                "name": "获奖/表彰",
                "name_en": "Awards",
                "description": "在该公司获得的奖项、表彰",
                "required": False,
                "file_types": ["pdf", "jpg", "png"],
                "has_form": False,
                "multiple": True,
                "tips": "如有获奖证书、表彰邮件等"
            }
        ]
    },
    "folder_4": {
        "name": "重大业绩证据",
        "name_en": "Major Achievements Evidence",
        "description": "展示您专业成就的证据材料",
        "order": 4,
        "items": [
            {
                "item_id": "contribution_form",
                "name": "原创贡献采集表",
                "name_en": "Original Contribution Form",
                "description": "独创、原创性重大贡献的详细描述",
                "required": True,
                "file_types": ["docx", "pdf"],
                "has_form": True,
                "form_type": "contribution",
                "tips": "详细描述您在行业内解决的问题和创新贡献"
            },
            {
                "item_id": "patents",
                "name": "专利证书",
                "name_en": "Patent Certificates",
                "description": "您参与的发明专利证书",
                "required": False,
                "file_types": ["pdf", "jpg", "png"],
                "has_form": False,
                "multiple": True,
                "tips": "提供专利证书扫描件，包括发明人页"
            },
            {
                "item_id": "publications",
                "name": "论文/出版物",
                "name_en": "Publications",
                "description": "发表的学术论文、技术文章等",
                "required": False,
                "file_types": ["pdf"],
                "has_form": False,
                "multiple": True,
                "tips": "提供论文PDF和引用数据截图"
            },
            {
                "item_id": "achievement_awards",
                "name": "成就奖项",
                "name_en": "Achievement Awards",
                "description": "行业奖项、竞赛获奖等证明",
                "required": False,
                "file_types": ["pdf", "jpg", "png"],
                "has_form": False,
                "multiple": True,
                "tips": "如Kaggle比赛名次、黑客松获奖等"
            },
            {
                "item_id": "opensource",
                "name": "开源项目证明",
                "name_en": "Open Source Contributions",
                "description": "开源项目贡献的截图和链接",
                "required": False,
                "file_types": ["pdf", "jpg", "png"],
                "has_form": True,
                "form_type": "opensource_links",
                "tips": "提供GitHub等平台的贡献统计截图"
            }
        ]
    },
    "folder_5": {
        "name": "项目证据",
        "name_en": "Project Evidence",
        "description": "重要项目的详细说明和证据（可添加多个项目）",
        "order": 5,
        "is_repeatable": True,  # 整个分类可重复添加
        "repeat_label": "项目",
        "min_count": 1,
        "max_count": 10,
        "items": [
            {
                "item_id": "project_name",
                "name": "项目名称",
                "name_en": "Project Name",
                "description": "项目基本信息",
                "required": True,
                "file_types": [],
                "has_form": True,
                "form_type": "project_basic",
                "tips": "填写项目名称、时间、您的角色"
            },
            {
                "item_id": "project_form",
                "name": "项目阐述表",
                "name_en": "Project Description Form",
                "description": "项目的详细描述和您的贡献",
                "required": True,
                "file_types": ["docx", "pdf"],
                "has_form": True,
                "form_type": "project",
                "tips": "详细描述项目背景、目标、您的贡献和成果"
            },
            {
                "item_id": "project_docs",
                "name": "项目证明材料",
                "name_en": "Project Documents",
                "description": "项目相关的文档、报告、截图等",
                "required": False,
                "file_types": ["pdf", "docx", "jpg", "png"],
                "has_form": False,
                "multiple": True,
                "tips": "如项目报告、上线公告、用户数据截图等"
            }
        ]
    },
    "folder_6": {
        "name": "推荐人材料",
        "name_en": "Recommender Documents",
        "description": "三位推荐人的原始信息材料（GTV签证要求至少3封推荐信，此处收集用于制作推荐信的原始素材）",
        "order": 6,
        "items": [
            # ===== 推荐人1 =====
            {
                "item_id": "recommender_1_resume",
                "name": "推荐人1-简历/履历",
                "name_en": "Recommender 1 - Resume/CV",
                "description": "第一位推荐人的个人简历或履历",
                "required": True,
                "file_types": ["pdf", "docx", "doc"],
                "has_form": False,
                "multiple": False,
                "tips": "推荐人的详细工作经历和教育背景"
            },
            {
                "item_id": "recommender_1_public_info",
                "name": "推荐人1-公开信息",
                "name_en": "Recommender 1 - Public Profile",
                "description": "推荐人的公开可查信息（LinkedIn、公司官网、学术主页等）",
                "required": True,
                "file_types": ["pdf", "jpg", "png"],
                "has_form": True,
                "form_type": "recommender_public_info",
                "multiple": True,
                "tips": "提供LinkedIn页面截图、公司官网介绍、学术主页等"
            },
            {
                "item_id": "recommender_1_relationship",
                "name": "推荐人1-关系说明",
                "name_en": "Recommender 1 - Relationship",
                "description": "申请人描述与推荐人的关系和合作经历",
                "required": True,
                "file_types": ["docx", "pdf"],
                "has_form": True,
                "form_type": "recommender_relationship",
                "tips": "详细说明如何认识、合作过程、推荐人对您的了解程度"
            },
            {
                "item_id": "recommender_1_contribution_form",
                "name": "推荐人1-杰出贡献采集表",
                "name_en": "Recommender 1 - Contribution Form",
                "description": "推荐人填写的个人杰出贡献采集表",
                "required": True,
                "file_types": ["docx", "pdf"],
                "has_form": True,
                "form_type": "recommender_contribution",
                "tips": "由推荐人填写，描述对申请人专业能力的评价"
            },
            # ===== 推荐人2 =====
            {
                "item_id": "recommender_2_resume",
                "name": "推荐人2-简历/履历",
                "name_en": "Recommender 2 - Resume/CV",
                "description": "第二位推荐人的个人简历或履历",
                "required": True,
                "file_types": ["pdf", "docx", "doc"],
                "has_form": False,
                "multiple": False,
                "tips": "推荐人的详细工作经历和教育背景"
            },
            {
                "item_id": "recommender_2_public_info",
                "name": "推荐人2-公开信息",
                "name_en": "Recommender 2 - Public Profile",
                "description": "推荐人的公开可查信息（LinkedIn、公司官网、学术主页等）",
                "required": True,
                "file_types": ["pdf", "jpg", "png"],
                "has_form": True,
                "form_type": "recommender_public_info",
                "multiple": True,
                "tips": "提供LinkedIn页面截图、公司官网介绍、学术主页等"
            },
            {
                "item_id": "recommender_2_relationship",
                "name": "推荐人2-关系说明",
                "name_en": "Recommender 2 - Relationship",
                "description": "申请人描述与推荐人的关系和合作经历",
                "required": True,
                "file_types": ["docx", "pdf"],
                "has_form": True,
                "form_type": "recommender_relationship",
                "tips": "详细说明如何认识、合作过程、推荐人对您的了解程度"
            },
            {
                "item_id": "recommender_2_contribution_form",
                "name": "推荐人2-杰出贡献采集表",
                "name_en": "Recommender 2 - Contribution Form",
                "description": "推荐人填写的个人杰出贡献采集表",
                "required": True,
                "file_types": ["docx", "pdf"],
                "has_form": True,
                "form_type": "recommender_contribution",
                "tips": "由推荐人填写，描述对申请人专业能力的评价"
            },
            # ===== 推荐人3 =====
            {
                "item_id": "recommender_3_resume",
                "name": "推荐人3-简历/履历",
                "name_en": "Recommender 3 - Resume/CV",
                "description": "第三位推荐人的个人简历或履历",
                "required": True,
                "file_types": ["pdf", "docx", "doc"],
                "has_form": False,
                "multiple": False,
                "tips": "推荐人的详细工作经历和教育背景"
            },
            {
                "item_id": "recommender_3_public_info",
                "name": "推荐人3-公开信息",
                "name_en": "Recommender 3 - Public Profile",
                "description": "推荐人的公开可查信息（LinkedIn、公司官网、学术主页等）",
                "required": True,
                "file_types": ["pdf", "jpg", "png"],
                "has_form": True,
                "form_type": "recommender_public_info",
                "multiple": True,
                "tips": "提供LinkedIn页面截图、公司官网介绍、学术主页等"
            },
            {
                "item_id": "recommender_3_relationship",
                "name": "推荐人3-关系说明",
                "name_en": "Recommender 3 - Relationship",
                "description": "申请人描述与推荐人的关系和合作经历",
                "required": True,
                "file_types": ["docx", "pdf"],
                "has_form": True,
                "form_type": "recommender_relationship",
                "tips": "详细说明如何认识、合作过程、推荐人对您的了解程度"
            },
            {
                "item_id": "recommender_3_contribution_form",
                "name": "推荐人3-杰出贡献采集表",
                "name_en": "Recommender 3 - Contribution Form",
                "description": "推荐人填写的个人杰出贡献采集表",
                "required": True,
                "file_types": ["docx", "pdf"],
                "has_form": True,
                "form_type": "recommender_contribution",
                "tips": "由推荐人填写，描述对申请人专业能力的评价"
            }
        ]
    }
}

# ==================== 从配置文件加载标签配置 ====================
def _load_material_categories():
    """从配置文件加载材料分类配置（如果存在）"""
    global MATERIAL_CATEGORIES
    config_path = os.path.join(os.path.dirname(__file__), 'material_categories.json')
    if os.path.exists(config_path):
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                loaded_categories = json.load(f)
            if loaded_categories and isinstance(loaded_categories, dict):
                MATERIAL_CATEGORIES = loaded_categories
                logger.info(f"从配置文件加载材料分类: {config_path}")
        except Exception as e:
            logger.warning(f"加载材料分类配置失败，使用默认配置: {e}")

# 模块加载时尝试加载配置
_load_material_categories()

# ==================== 采集表单模板 ====================

FORM_TEMPLATES = {
    "employment_info": {
        "title": "就职信息采集表",
        "description": "请详细填写您当前的工作信息",
        "fields": [
            {"name": "company_name", "label": "公司名称", "type": "text", "required": True},
            {"name": "employment_period", "label": "就职时间段", "type": "text", "required": True, "placeholder": "如：2021.04-至今"},
            {"name": "position", "label": "职位", "type": "text", "required": True},
            {"name": "is_legal_person", "label": "是否法人", "type": "select", "options": ["是", "否"], "required": True},
            {"name": "is_shareholder", "label": "是否股东", "type": "select", "options": ["是", "否"], "required": True},
            {"name": "responsibilities", "label": "工作职责（详细描述）", "type": "textarea", "required": True},
            {"name": "project_1", "label": "项目一（名称和成果）", "type": "textarea", "required": False},
            {"name": "project_2", "label": "项目二（名称和成果）", "type": "textarea", "required": False},
            {"name": "project_3", "label": "项目三（名称和成果）", "type": "textarea", "required": False},
            {"name": "major_clients", "label": "合作过的重要客户", "type": "textarea", "required": False},
            {"name": "income", "label": "缴税收入（如有期权股权请额外说明）", "type": "textarea", "required": False}
        ]
    },
    "prev_company_basic": {
        "title": "过往公司基本信息",
        "description": "填写过往工作单位的基本信息",
        "fields": [
            {"name": "company_name", "label": "公司名称", "type": "text", "required": True},
            {"name": "company_name_en", "label": "公司英文名", "type": "text", "required": False},
            {"name": "employment_period", "label": "就职时间段", "type": "text", "required": True, "placeholder": "如：2018.06-2021.03"},
            {"name": "position", "label": "职位", "type": "text", "required": True},
            {"name": "position_en", "label": "职位英文", "type": "text", "required": False},
            {"name": "company_intro", "label": "公司简介", "type": "textarea", "required": False, "placeholder": "公司规模、行业地位、主营业务等"}
        ]
    },
    "prev_employment": {
        "title": "过往就职详细信息",
        "description": "详细描述在该公司的工作内容和成就",
        "fields": [
            {"name": "responsibilities", "label": "工作职责（详细描述）", "type": "textarea", "required": True},
            {"name": "key_projects", "label": "主要参与/负责的项目", "type": "textarea", "required": True, "placeholder": "列举主要项目及您的角色"},
            {"name": "achievements", "label": "主要成就和贡献", "type": "textarea", "required": True, "placeholder": "量化的成果、获得的认可等"},
            {"name": "awards", "label": "获得的奖项/表彰", "type": "textarea", "required": False},
            {"name": "skills_gained", "label": "获得的技能/成长", "type": "textarea", "required": False},
            {"name": "reason_leave", "label": "离职原因", "type": "text", "required": False}
        ]
    },
    "project_basic": {
        "title": "项目基本信息",
        "description": "填写项目的基本信息",
        "fields": [
            {"name": "project_name", "label": "项目名称", "type": "text", "required": True},
            {"name": "project_name_en", "label": "项目英文名", "type": "text", "required": False},
            {"name": "project_period", "label": "项目时间", "type": "text", "required": True, "placeholder": "如：2020.01-2021.06"},
            {"name": "your_role", "label": "您的角色", "type": "text", "required": True, "placeholder": "如：技术负责人、架构师、项目经理"},
            {"name": "team_size", "label": "团队规模", "type": "text", "required": False},
            {"name": "company", "label": "所属公司", "type": "text", "required": False}
        ]
    },
    "contribution": {
        "title": "独创原创性重大贡献采集表",
        "description": "请描述您在专业领域内的重要贡献",
        "fields": [
            {"name": "contribution_title", "label": "贡献名称", "type": "text", "required": True, "placeholder": "独创***技术/运营模式/营销设想/理论等"},
            {"name": "industry_problems", "label": "您所在领域存在的行业主要问题", "type": "textarea", "required": True},
            {"name": "problems_solved", "label": "您解决了哪些行业问题？", "type": "textarea", "required": True},
            {"name": "how_solved", "label": "您通过哪些原创贡献解决了以上问题？", "type": "textarea", "required": True},
            {"name": "impact", "label": "您的原创对于整个行业带来了什么影响？", "type": "textarea", "required": True, "placeholder": "短期冲击/长期影响/意义重大的水平"},
            {"name": "scope", "label": "您的原创的适用范围、区域", "type": "textarea", "required": False},
            {"name": "economic_benefit", "label": "带来的经济效益", "type": "textarea", "required": False},
            {"name": "social_benefit", "label": "带来的社会效益", "type": "textarea", "required": False},
            {"name": "peer_impact", "label": "对同行的改变", "type": "textarea", "required": False},
            {"name": "expert_praise", "label": "同领域专家、领军人对重大贡献的赞扬、证词、引用等", "type": "textarea", "required": False},
            {"name": "uk_benefit", "label": "您将给英国带来什么利益？", "type": "textarea", "required": True}
        ]
    },
    "project": {
        "title": "大型项目阐述表",
        "description": "请详细描述一个您主导或深度参与的重要项目",
        "fields": [
            {"name": "project_name", "label": "项目名称", "type": "text", "required": True},
            {"name": "project_period", "label": "项目时间", "type": "text", "required": True},
            {"name": "your_role", "label": "您的角色", "type": "text", "required": True},
            {"name": "team_size", "label": "团队规模", "type": "text", "required": False},
            {"name": "project_background", "label": "项目背景", "type": "textarea", "required": True},
            {"name": "project_goals", "label": "项目目标", "type": "textarea", "required": True},
            {"name": "your_contribution", "label": "您的具体贡献", "type": "textarea", "required": True},
            {"name": "challenges", "label": "遇到的挑战及解决方案", "type": "textarea", "required": False},
            {"name": "results", "label": "项目成果（量化数据）", "type": "textarea", "required": True},
            {"name": "impact", "label": "项目影响力", "type": "textarea", "required": False}
        ]
    },
    "media_links": {
        "title": "媒体报道链接",
        "description": "请提供关于您或公司的媒体报道链接",
        "fields": [
            {"name": "link_1", "label": "媒体链接1", "type": "url", "required": False},
            {"name": "link_1_desc", "label": "报道简述1", "type": "text", "required": False},
            {"name": "link_2", "label": "媒体链接2", "type": "url", "required": False},
            {"name": "link_2_desc", "label": "报道简述2", "type": "text", "required": False},
            {"name": "link_3", "label": "媒体链接3", "type": "url", "required": False},
            {"name": "link_3_desc", "label": "报道简述3", "type": "text", "required": False}
        ]
    },
    "opensource_links": {
        "title": "开源项目贡献",
        "description": "请提供您的开源项目信息",
        "fields": [
            {"name": "github_url", "label": "GitHub主页", "type": "url", "required": False},
            {"name": "stars_count", "label": "总Star数", "type": "text", "required": False},
            {"name": "project_1_url", "label": "项目1链接", "type": "url", "required": False},
            {"name": "project_1_desc", "label": "项目1描述", "type": "text", "required": False},
            {"name": "project_2_url", "label": "项目2链接", "type": "url", "required": False},
            {"name": "project_2_desc", "label": "项目2描述", "type": "text", "required": False},
            {"name": "contributions", "label": "主要贡献描述", "type": "textarea", "required": False}
        ]
    },
    "recommender_public_info": {
        "title": "推荐人公开信息采集",
        "description": "请收集并填写推荐人的公开可查信息",
        "fields": [
            {"name": "recommender_name", "label": "推荐人姓名", "type": "text", "required": True},
            {"name": "recommender_name_en", "label": "推荐人英文名", "type": "text", "required": True},
            {"name": "current_title", "label": "当前职位", "type": "text", "required": True},
            {"name": "current_company", "label": "当前公司/机构", "type": "text", "required": True},
            {"name": "linkedin_url", "label": "LinkedIn链接", "type": "url", "required": False},
            {"name": "company_page_url", "label": "公司官网个人介绍页", "type": "url", "required": False},
            {"name": "academic_page_url", "label": "学术主页/Google Scholar", "type": "url", "required": False},
            {"name": "other_public_url", "label": "其他公开信息链接", "type": "url", "required": False},
            {"name": "public_achievements", "label": "公开可查的主要成就", "type": "textarea", "required": True, "placeholder": "如：发表论文数、专利数、获奖情况、媒体报道等"},
            {"name": "industry_reputation", "label": "行业地位和声誉", "type": "textarea", "required": True, "placeholder": "推荐人在行业内的知名度、影响力"}
        ]
    },
    "recommender_relationship": {
        "title": "与推荐人关系说明",
        "description": "申请人填写：详细说明与推荐人的关系和合作经历",
        "fields": [
            {"name": "recommender_name", "label": "推荐人姓名", "type": "text", "required": True},
            {"name": "relationship_type", "label": "关系类型", "type": "select", "options": ["直属上级", "公司高管", "合作伙伴", "客户", "学术导师", "行业专家", "投资人", "其他"], "required": True},
            {"name": "how_met", "label": "如何认识推荐人", "type": "textarea", "required": True, "placeholder": "描述在什么场合、通过什么方式认识的"},
            {"name": "know_duration", "label": "认识时长", "type": "text", "required": True, "placeholder": "如：5年"},
            {"name": "collaboration_details", "label": "合作经历详情", "type": "textarea", "required": True, "placeholder": "详细描述与推荐人的合作项目、工作关系、共同成就等"},
            {"name": "recommender_witness", "label": "推荐人见证了您的哪些成就", "type": "textarea", "required": True, "placeholder": "推荐人亲眼见证或参与的您的专业成就"},
            {"name": "why_choose", "label": "为什么选择此人作为推荐人", "type": "textarea", "required": True, "placeholder": "推荐人的行业地位、对您工作的了解程度等"},
            {"name": "contact_email", "label": "推荐人联系邮箱", "type": "text", "required": True},
            {"name": "contact_phone", "label": "推荐人联系电话", "type": "text", "required": False}
        ]
    },
    "recommender_contribution": {
        "title": "个人杰出贡献采集表（推荐人填写）",
        "description": "请推荐人根据对申请人的了解，详细描述申请人的专业贡献和行业影响力",
        "fields": [
            {"name": "industry_background", "label": "行业背景：目前该行业存在的痛点和主要问题是什么？", "type": "textarea", "required": True},
            {"name": "applicant_role", "label": "申请人在解决这些痛点问题上发挥了什么样的作用？", "type": "textarea", "required": True},
            {"name": "company_strength", "label": "申请人所在公司实力如何？在行业里处于什么水准？", "type": "textarea", "required": True},
            {"name": "industry_standards", "label": "该行业对于优秀人才是否有通用的评定标准？", "type": "textarea", "required": False},
            {"name": "applicant_level", "label": "申请人在公司内部的杰出性处于什么水平？", "type": "textarea", "required": True},
            {"name": "personal_achievement", "label": "申请人的个人成就/原创贡献，对整个行业和公司带来了什么影响？", "type": "textarea", "required": True, "placeholder": "短期/长期影响，经济价值和社会价值"},
            {"name": "peer_impact", "label": "申请人的贡献对同行产生了什么影响？", "type": "textarea", "required": False},
            {"name": "third_party_evaluation", "label": "是否有第三方媒体/个人对申请人能力的评价？", "type": "textarea", "required": False},
            {"name": "uk_industry_status", "label": "对申请人所在行业在英国现状的评价", "type": "textarea", "required": True},
            {"name": "uk_attraction", "label": "英国哪些方面吸引申请人？为什么考虑英国发展？", "type": "textarea", "required": True},
            {"name": "uk_plan", "label": "申请人拿到身份后的计划（创业/就业、定居城市等）", "type": "textarea", "required": True},
            {"name": "startup_plan", "label": "如选择创业，请描述创业思路和周期", "type": "textarea", "required": False},
            {"name": "uk_connections", "label": "申请人是否已与英国行业资源对接？请提供细节", "type": "textarea", "required": False}
        ]
    }
}


class RawMaterialManager:
    """原始材料收集管理器"""
    
    def __init__(self, db_path: str = None, upload_folder: str = None):
        """
        初始化管理器
        
        Args:
            db_path: 数据库文件路径
            upload_folder: 文件上传目录
        """
        self.db_path = db_path or os.getenv("COPYWRITING_DB_PATH", "copywriting.db")
        self.upload_folder = upload_folder or os.getenv("UPLOAD_FOLDER", "./uploads")
        Path(self.upload_folder).mkdir(parents=True, exist_ok=True)
        self._init_tables()
        logger.info(f"原始材料管理器初始化完成")
    
    @contextmanager
    def _get_connection(self):
        """获取数据库连接"""
        conn = sqlite3.connect(self.db_path)
        conn.row_factory = sqlite3.Row
        try:
            yield conn
            conn.commit()
        except Exception as e:
            conn.rollback()
            raise e
        finally:
            conn.close()
    
    def _init_tables(self):
        """初始化数据库表"""
        try:
            with self._get_connection() as conn:
                cursor = conn.cursor()
                
                # 材料收集状态表
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS material_collection (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        project_id TEXT NOT NULL,
                        category_id TEXT NOT NULL,
                        item_id TEXT NOT NULL,
                        status TEXT DEFAULT 'pending',
                        file_path TEXT,
                        file_name TEXT,
                        file_size INTEGER DEFAULT 0,
                        file_type TEXT,
                        form_data TEXT,
                        notes TEXT,
                        collected_at TIMESTAMP,
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        UNIQUE(project_id, category_id, item_id)
                    )
                ''')
                
                # 多文件上传记录表（支持同一item多个文件）
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS material_files (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        project_id TEXT NOT NULL,
                        category_id TEXT NOT NULL,
                        item_id TEXT NOT NULL,
                        file_name TEXT NOT NULL,
                        file_path TEXT NOT NULL,
                        file_size INTEGER DEFAULT 0,
                        file_type TEXT,
                        description TEXT,
                        uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                    )
                ''')
                
                # 采集表单数据表
                cursor.execute('''
                    CREATE TABLE IF NOT EXISTS collection_forms (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        project_id TEXT NOT NULL,
                        form_type TEXT NOT NULL,
                        form_index INTEGER DEFAULT 0,
                        form_data TEXT NOT NULL,
                        status TEXT DEFAULT 'draft',
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        UNIQUE(project_id, form_type, form_index)
                    )
                ''')
                
                # 创建索引
                cursor.execute('CREATE INDEX IF NOT EXISTS idx_material_collection_project ON material_collection (project_id)')
                cursor.execute('CREATE INDEX IF NOT EXISTS idx_material_files_project ON material_files (project_id)')
                cursor.execute('CREATE INDEX IF NOT EXISTS idx_collection_forms_project ON collection_forms (project_id)')
                
                conn.commit()
                logger.info("原始材料表结构初始化完成")
                
        except Exception as e:
            logger.error(f"初始化原始材料表失败: {e}")
            raise
    
    # ==================== 材料分类和模板 ====================
    
    def get_material_categories(self) -> Dict[str, Any]:
        """获取材料分类结构"""
        return {
            "success": True,
            "data": MATERIAL_CATEGORIES
        }
    
    def get_form_template(self, form_type: str) -> Dict[str, Any]:
        """获取采集表单模板"""
        if form_type in FORM_TEMPLATES:
            return {
                "success": True,
                "data": FORM_TEMPLATES[form_type]
            }
        return {
            "success": False,
            "error": f"未知的表单类型: {form_type}"
        }
    
    def get_all_form_templates(self) -> Dict[str, Any]:
        """获取所有表单模板"""
        return {
            "success": True,
            "data": FORM_TEMPLATES
        }
    
    # ==================== 项目材料初始化 ====================
    
    def init_project_materials(self, project_id: str) -> Dict[str, Any]:
        """为项目初始化材料收集清单"""
        try:
            with self._get_connection() as conn:
                cursor = conn.cursor()
                
                # 为每个材料项创建记录
                for cat_id, category in MATERIAL_CATEGORIES.items():
                    for item in category["items"]:
                        cursor.execute('''
                            INSERT OR IGNORE INTO material_collection 
                            (project_id, category_id, item_id, status)
                            VALUES (?, ?, ?, 'pending')
                        ''', (project_id, cat_id, item["item_id"]))
                
                conn.commit()
                logger.info(f"项目 {project_id} 材料清单初始化完成")
                return {"success": True}
                
        except Exception as e:
            logger.error(f"初始化项目材料清单失败: {e}")
            return {"success": False, "error": str(e)}
    
    # ==================== 材料收集状态 ====================
    
    def get_collection_status(self, project_id: str) -> Dict[str, Any]:
        """获取项目材料收集状态"""
        try:
            # 先确保项目已初始化
            self.init_project_materials(project_id)
            
            with self._get_connection() as conn:
                cursor = conn.cursor()
                
                # 获取收集状态
                cursor.execute('''
                    SELECT * FROM material_collection WHERE project_id = ?
                ''', (project_id,))
                
                collection_map = {}
                for row in cursor.fetchall():
                    key = f"{row['category_id']}_{row['item_id']}"
                    collection_map[key] = {
                        "status": row["status"],
                        "file_path": row["file_path"],
                        "file_name": row["file_name"],
                        "collected_at": row["collected_at"],
                        "notes": row["notes"]
                    }
                
                # 获取多文件记录
                cursor.execute('''
                    SELECT * FROM material_files WHERE project_id = ? ORDER BY uploaded_at
                ''', (project_id,))
                
                files_map = {}
                for row in cursor.fetchall():
                    key = f"{row['category_id']}_{row['item_id']}"
                    if key not in files_map:
                        files_map[key] = []
                    files_map[key].append({
                        "id": row["id"],
                        "file_name": row["file_name"],
                        "file_path": row["file_path"],
                        "file_size": row["file_size"],
                        "file_type": row["file_type"],
                        "description": row["description"],
                        "uploaded_at": row["uploaded_at"]
                    })
                
                # 构建完整的状态数据
                result = {}
                total_items = 0
                collected_items = 0
                required_items = 0
                required_collected = 0
                
                for cat_id, category in MATERIAL_CATEGORIES.items():
                    cat_data = {
                        "name": category["name"],
                        "name_en": category["name_en"],
                        "description": category["description"],
                        "order": category["order"],
                        "items": []
                    }
                    
                    for item in category["items"]:
                        key = f"{cat_id}_{item['item_id']}"
                        collection = collection_map.get(key, {"status": "pending"})
                        files = files_map.get(key, [])
                        
                        item_data = {
                            **item,
                            "status": collection["status"],
                            "file_name": collection.get("file_name"),
                            "collected_at": collection.get("collected_at"),
                            "notes": collection.get("notes"),
                            "files": files
                        }
                        
                        cat_data["items"].append(item_data)
                        
                        total_items += 1
                        if collection["status"] == "collected":
                            collected_items += 1
                        
                        if item.get("required"):
                            required_items += 1
                            if collection["status"] == "collected":
                                required_collected += 1
                    
                    result[cat_id] = cat_data
                
                # 计算进度
                progress = {
                    "total_items": total_items,
                    "collected_items": collected_items,
                    "required_items": required_items,
                    "required_collected": required_collected,
                    "overall_progress": round(collected_items / total_items * 100) if total_items > 0 else 0,
                    "required_progress": round(required_collected / required_items * 100) if required_items > 0 else 0
                }
                
                return {
                    "success": True,
                    "data": {
                        "categories": result,
                        "progress": progress
                    }
                }
                
        except Exception as e:
            logger.error(f"获取材料收集状态失败: {e}")
            return {"success": False, "error": str(e)}
    
    # ==================== 文件上传 ====================
    
    def upload_material(self, project_id: str, category_id: str, item_id: str,
                       file_path: str, file_name: str, file_size: int = 0,
                       file_type: str = None, description: str = None) -> Dict[str, Any]:
        """上传材料文件"""
        try:
            with self._get_connection() as conn:
                cursor = conn.cursor()
                
                # 检查是否支持多文件
                item_info = None
                for cat_id, cat in MATERIAL_CATEGORIES.items():
                    if cat_id == category_id:
                        for item in cat["items"]:
                            if item["item_id"] == item_id:
                                item_info = item
                                break
                
                if not item_info:
                    return {"success": False, "error": "未知的材料项"}
                
                # 记录文件
                cursor.execute('''
                    INSERT INTO material_files 
                    (project_id, category_id, item_id, file_name, file_path, file_size, file_type, description)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''', (project_id, category_id, item_id, file_name, file_path, file_size, file_type, description))
                
                file_id = cursor.lastrowid
                
                # 更新收集状态
                cursor.execute('''
                    INSERT OR REPLACE INTO material_collection 
                    (project_id, category_id, item_id, status, file_path, file_name, file_size, file_type, collected_at, updated_at)
                    VALUES (?, ?, ?, 'collected', ?, ?, ?, ?, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
                ''', (project_id, category_id, item_id, file_path, file_name, file_size, file_type))
                
                conn.commit()
                logger.info(f"材料上传成功: {project_id}/{category_id}/{item_id}")
                
                return {
                    "success": True,
                    "file_id": file_id,
                    "message": "文件上传成功"
                }
                
        except Exception as e:
            logger.error(f"上传材料失败: {e}")
            return {"success": False, "error": str(e)}
    
    def delete_material_file(self, file_id: int) -> Dict[str, Any]:
        """删除材料文件"""
        try:
            with self._get_connection() as conn:
                cursor = conn.cursor()
                
                # 获取文件信息
                cursor.execute('SELECT * FROM material_files WHERE id = ?', (file_id,))
                row = cursor.fetchone()
                
                if not row:
                    return {"success": False, "error": "文件不存在"}
                
                project_id = row["project_id"]
                category_id = row["category_id"]
                item_id = row["item_id"]
                file_path = row["file_path"]
                
                # 删除数据库记录
                cursor.execute('DELETE FROM material_files WHERE id = ?', (file_id,))
                
                # 检查是否还有其他文件
                cursor.execute('''
                    SELECT COUNT(*) as count FROM material_files 
                    WHERE project_id = ? AND category_id = ? AND item_id = ?
                ''', (project_id, category_id, item_id))
                
                remaining = cursor.fetchone()["count"]
                
                # 如果没有其他文件，更新状态为pending
                if remaining == 0:
                    cursor.execute('''
                        UPDATE material_collection 
                        SET status = 'pending', file_path = NULL, file_name = NULL, collected_at = NULL
                        WHERE project_id = ? AND category_id = ? AND item_id = ?
                    ''', (project_id, category_id, item_id))
                
                conn.commit()
                
                # 删除物理文件
                if file_path and os.path.exists(file_path):
                    os.remove(file_path)
                
                return {"success": True, "message": "文件删除成功"}
                
        except Exception as e:
            logger.error(f"删除材料文件失败: {e}")
            return {"success": False, "error": str(e)}
    
    def update_material_tags(self, file_id: int, category_id: str, item_id: str) -> Dict[str, Any]:
        """更新材料文件的分类标签"""
        try:
            with self._get_connection() as conn:
                cursor = conn.cursor()
                
                # 获取文件信息
                cursor.execute('SELECT * FROM material_files WHERE id = ?', (file_id,))
                row = cursor.fetchone()
                
                if not row:
                    return {"success": False, "error": "文件不存在"}
                
                old_category_id = row["category_id"]
                old_item_id = row["item_id"]
                project_id = row["project_id"]
                file_name = row["file_name"]
                file_size = row["file_size"]
                file_type = row["file_type"]
                file_path = row["file_path"]
                
                # 更新文件的分类
                cursor.execute('''
                    UPDATE material_files 
                    SET category_id = ?, item_id = ?
                    WHERE id = ?
                ''', (category_id, item_id, file_id))
                
                # 更新旧分类的状态（如果没有其他文件了）
                cursor.execute('''
                    SELECT COUNT(*) as count FROM material_files 
                    WHERE project_id = ? AND category_id = ? AND item_id = ?
                ''', (project_id, old_category_id, old_item_id))
                
                remaining = cursor.fetchone()["count"]
                
                if remaining == 0:
                    cursor.execute('''
                        UPDATE material_collection 
                        SET status = 'pending', file_path = NULL, file_name = NULL, collected_at = NULL
                        WHERE project_id = ? AND category_id = ? AND item_id = ?
                    ''', (project_id, old_category_id, old_item_id))
                
                # 更新新分类的状态
                cursor.execute('''
                    INSERT OR REPLACE INTO material_collection 
                    (project_id, category_id, item_id, status, file_path, file_name, file_size, file_type, collected_at, updated_at)
                    VALUES (?, ?, ?, 'collected', ?, ?, ?, ?, datetime('now'), datetime('now'))
                ''', (project_id, category_id, item_id, file_path, file_name, file_size, file_type))
                
                conn.commit()
                
                logger.info(f"材料标签更新成功: file_id={file_id}, {old_category_id}/{old_item_id} -> {category_id}/{item_id}")
                return {"success": True, "message": "标签更新成功"}
                
        except Exception as e:
            logger.error(f"更新材料标签失败: {e}")
            return {"success": False, "error": str(e)}
    
    # ==================== 表单数据 ====================
    
    def save_form_data(self, project_id: str, form_type: str, form_data: Dict,
                      form_index: int = 0) -> Dict[str, Any]:
        """保存采集表单数据"""
        try:
            with self._get_connection() as conn:
                cursor = conn.cursor()
                
                cursor.execute('''
                    INSERT OR REPLACE INTO collection_forms 
                    (project_id, form_type, form_index, form_data, status, updated_at)
                    VALUES (?, ?, ?, ?, 'completed', CURRENT_TIMESTAMP)
                ''', (project_id, form_type, form_index, json.dumps(form_data, ensure_ascii=False)))
                
                form_id = cursor.lastrowid
                
                # 更新对应材料项状态
                # 找到使用这个form_type的材料项
                for cat_id, cat in MATERIAL_CATEGORIES.items():
                    for item in cat["items"]:
                        if item.get("form_type") == form_type:
                            cursor.execute('''
                                INSERT OR REPLACE INTO material_collection 
                                (project_id, category_id, item_id, status, form_data, collected_at, updated_at)
                                VALUES (?, ?, ?, 'collected', ?, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
                            ''', (project_id, cat_id, item["item_id"], json.dumps(form_data, ensure_ascii=False)))
                
                conn.commit()
                logger.info(f"表单数据保存成功: {project_id}/{form_type}")
                
                return {
                    "success": True,
                    "form_id": form_id,
                    "message": "表单保存成功"
                }
                
        except Exception as e:
            logger.error(f"保存表单数据失败: {e}")
            return {"success": False, "error": str(e)}
    
    def get_form_data(self, project_id: str, form_type: str, 
                     form_index: int = 0) -> Dict[str, Any]:
        """获取表单数据"""
        try:
            with self._get_connection() as conn:
                cursor = conn.cursor()
                
                cursor.execute('''
                    SELECT * FROM collection_forms 
                    WHERE project_id = ? AND form_type = ? AND form_index = ?
                ''', (project_id, form_type, form_index))
                
                row = cursor.fetchone()
                if row:
                    return {
                        "success": True,
                        "data": {
                            "id": row["id"],
                            "form_type": row["form_type"],
                            "form_index": row["form_index"],
                            "form_data": json.loads(row["form_data"]),
                            "status": row["status"],
                            "created_at": row["created_at"],
                            "updated_at": row["updated_at"]
                        }
                    }
                
                return {"success": True, "data": None}
                
        except Exception as e:
            logger.error(f"获取表单数据失败: {e}")
            return {"success": False, "error": str(e)}
    
    def get_all_forms(self, project_id: str) -> Dict[str, Any]:
        """获取项目所有表单数据"""
        try:
            with self._get_connection() as conn:
                cursor = conn.cursor()
                
                cursor.execute('''
                    SELECT * FROM collection_forms 
                    WHERE project_id = ? ORDER BY form_type, form_index
                ''', (project_id,))
                
                forms = {}
                for row in cursor.fetchall():
                    form_type = row["form_type"]
                    if form_type not in forms:
                        forms[form_type] = []
                    
                    forms[form_type].append({
                        "id": row["id"],
                        "form_index": row["form_index"],
                        "form_data": json.loads(row["form_data"]),
                        "status": row["status"],
                        "updated_at": row["updated_at"]
                    })
                
                return {"success": True, "data": forms}
                
        except Exception as e:
            logger.error(f"获取所有表单失败: {e}")
            return {"success": False, "error": str(e)}
    
    # ==================== 材料完整性检查 ====================
    
    def check_completeness(self, project_id: str) -> Dict[str, Any]:
        """检查材料完整性"""
        try:
            status_result = self.get_collection_status(project_id)
            if not status_result["success"]:
                return status_result
            
            categories = status_result["data"]["categories"]
            progress = status_result["data"]["progress"]
            
            missing_required = []
            missing_optional = []
            collected = []
            
            for cat_id, category in categories.items():
                for item in category["items"]:
                    item_info = {
                        "category": category["name"],
                        "category_id": cat_id,
                        "item_id": item["item_id"],
                        "name": item["name"],
                        "required": item.get("required", False)
                    }
                    
                    if item["status"] == "collected":
                        collected.append(item_info)
                    elif item.get("required"):
                        missing_required.append(item_info)
                    else:
                        missing_optional.append(item_info)
            
            is_complete = len(missing_required) == 0
            
            return {
                "success": True,
                "data": {
                    "is_complete": is_complete,
                    "progress": progress,
                    "missing_required": missing_required,
                    "missing_optional": missing_optional,
                    "collected": collected,
                    "summary": {
                        "required_missing_count": len(missing_required),
                        "optional_missing_count": len(missing_optional),
                        "collected_count": len(collected)
                    }
                }
            }
            
        except Exception as e:
            logger.error(f"检查材料完整性失败: {e}")
            return {"success": False, "error": str(e)}
    
    # ==================== 导出材料清单 ====================
    
    def export_checklist(self, project_id: str) -> Dict[str, Any]:
        """导出材料收集清单（用于打印或发送给客户）"""
        try:
            status_result = self.get_collection_status(project_id)
            if not status_result["success"]:
                return status_result
            
            categories = status_result["data"]["categories"]
            
            # 生成Markdown格式的清单
            lines = ["# GTV签证申请材料收集清单\n"]
            lines.append(f"项目ID: {project_id}\n")
            lines.append(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
            lines.append("---\n")
            
            for cat_id, category in sorted(categories.items(), key=lambda x: x[1]["order"]):
                lines.append(f"\n## {category['name']} ({category['name_en']})\n")
                lines.append(f"{category['description']}\n")
                
                for item in category["items"]:
                    status_icon = "✅" if item["status"] == "collected" else "⬜"
                    required_mark = " *必填*" if item.get("required") else ""
                    
                    lines.append(f"\n### {status_icon} {item['name']}{required_mark}")
                    lines.append(f"英文名: {item['name_en']}")
                    lines.append(f"说明: {item['description']}")
                    
                    if item.get("tips"):
                        lines.append(f"提示: {item['tips']}")
                    
                    if item.get("file_types"):
                        lines.append(f"支持格式: {', '.join(item['file_types'])}")
                    
                    if item["status"] == "collected":
                        lines.append(f"状态: 已收集 ({item.get('collected_at', '')})")
                        if item.get("files"):
                            lines.append("已上传文件:")
                            for f in item["files"]:
                                lines.append(f"  - {f['file_name']}")
            
            checklist_content = "\n".join(lines)
            
            return {
                "success": True,
                "data": {
                    "content": checklist_content,
                    "format": "markdown"
                }
            }
            
        except Exception as e:
            logger.error(f"导出材料清单失败: {e}")
            return {"success": False, "error": str(e)}
    
    # ==================== 生成可打印的Word文档 ====================
    
    def generate_checklist_document(self, project_id: str, client_name: str = None,
                                   output_dir: str = None) -> Dict[str, Any]:
        """生成可打印的材料收集清单Word文档"""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx未安装，无法生成Word文档"}
        
        try:
            # 获取当前收集状态
            status_result = self.get_collection_status(project_id)
            categories = status_result.get("data", {}).get("categories", MATERIAL_CATEGORIES) if status_result.get("success") else MATERIAL_CATEGORIES
            
            # 创建文档
            doc = Document()
            
            # 设置页面边距
            sections = doc.sections
            for section in sections:
                section.top_margin = Cm(2)
                section.bottom_margin = Cm(2)
                section.left_margin = Cm(2.5)
                section.right_margin = Cm(2.5)
            
            # 标题
            title = doc.add_heading('GTV签证申请材料收集清单', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 客户信息
            if client_name:
                info_para = doc.add_paragraph()
                info_para.add_run(f'客户姓名：{client_name}').bold = True
                info_para.add_run(f'          生成日期：{datetime.now().strftime("%Y年%m月%d日")}')
                info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            
            # 说明
            intro = doc.add_paragraph()
            intro.add_run('使用说明：').bold = True
            doc.add_paragraph('1. 请按照清单逐项准备材料，在"已收集"列打勾 ☑')
            doc.add_paragraph('2. 标注"必填"的材料为申请必需，请务必提供')
            doc.add_paragraph('3. 需要填写采集表的材料，请下载对应模板填写后上传')
            doc.add_paragraph('4. 准备完成后，请将所有材料打包发送给您的顾问')
            
            doc.add_paragraph()
            
            # 遍历分类
            sorted_cats = sorted(
                (categories if isinstance(categories, dict) else MATERIAL_CATEGORIES).items(), 
                key=lambda x: x[1].get("order", 0)
            )
            
            for cat_id, category in sorted_cats:
                # 分类标题
                cat_heading = doc.add_heading(f'{category["name"]} ({category["name_en"]})', level=1)
                doc.add_paragraph(category["description"])
                
                # 创建表格
                items = category.get("items", [])
                if not items:
                    continue
                
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # 表头
                header_cells = table.rows[0].cells
                headers = ['序号', '材料名称', '说明/要求', '是否必填', '已收集']
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                    header_cells[i].paragraphs[0].runs[0].bold = True
                    header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 设置列宽
                widths = [Cm(1), Cm(3.5), Cm(8), Cm(2), Cm(2)]
                for i, width in enumerate(widths):
                    for cell in table.columns[i].cells:
                        cell.width = width
                
                # 填充数据
                for idx, item in enumerate(items, 1):
                    row = table.add_row().cells
                    row[0].text = str(idx)
                    row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # 材料名称
                    name_para = row[1].paragraphs[0]
                    name_para.add_run(item["name"]).bold = True
                    name_para.add_run(f'\n({item["name_en"]})')
                    
                    # 说明
                    desc_text = item["description"]
                    if item.get("tips"):
                        desc_text += f'\n💡 {item["tips"]}'
                    if item.get("file_types"):
                        desc_text += f'\n📎 格式: {", ".join(item["file_types"])}'
                    if item.get("has_form"):
                        desc_text += f'\n📝 需填写采集表'
                    if item.get("generated"):
                        desc_text += f'\n✨ 顾问协助准备'
                    row[2].text = desc_text
                    
                    # 必填
                    row[3].text = '必填' if item.get("required") else '选填'
                    row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if item.get("required"):
                        row[3].paragraphs[0].runs[0].bold = True
                    
                    # 已收集（勾选框）
                    status = item.get("status", "pending")
                    row[4].text = '☑' if status == 'collected' else '☐'
                    row[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                doc.add_paragraph()
            
            # 附录：采集表清单
            doc.add_page_break()
            doc.add_heading('附录：采集表模板说明', level=1)
            doc.add_paragraph('以下采集表需要您下载填写后上传：')
            doc.add_paragraph()
            
            form_list = [
                ('就职信息采集表', 'employment_info', '填写当前工作的详细信息'),
                ('过往就职信息表', 'prev_employment', '填写每段过往工作经历'),
                ('原创贡献采集表', 'contribution', '描述您的独创性贡献'),
                ('大型项目阐述表', 'project', '描述您主导的重要项目'),
            ]
            
            form_table = doc.add_table(rows=1, cols=3)
            form_table.style = 'Table Grid'
            
            header_cells = form_table.rows[0].cells
            headers = ['表单名称', '用途说明', '备注']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
            
            for form_name, form_type, desc in form_list:
                row = form_table.add_row().cells
                row[0].text = form_name
                row[1].text = desc
                row[2].text = '请向顾问索取模板'
            
            # 保存文档
            output_dir = output_dir or self.upload_folder
            Path(output_dir).mkdir(parents=True, exist_ok=True)
            
            safe_name = client_name.replace('/', '_').replace('\\', '_') if client_name else project_id
            filename = f'GTV材料清单_{safe_name}_{datetime.now().strftime("%Y%m%d")}.docx'
            file_path = os.path.join(output_dir, filename)
            
            doc.save(file_path)
            logger.info(f"材料清单文档生成成功: {file_path}")
            
            return {
                "success": True,
                "data": {
                    "file_path": file_path,
                    "filename": filename,
                    "format": "docx"
                }
            }
            
        except Exception as e:
            logger.error(f"生成材料清单文档失败: {e}")
            return {"success": False, "error": str(e)}
    
    def generate_form_template(self, form_type: str, output_dir: str = None) -> Dict[str, Any]:
        """生成单个采集表模板Word文档"""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx未安装"}
        
        if form_type not in FORM_TEMPLATES:
            return {"success": False, "error": f"未知的表单类型: {form_type}"}
        
        try:
            template = FORM_TEMPLATES[form_type]
            
            doc = Document()
            
            # 标题
            title = doc.add_heading(template["title"], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 说明
            doc.add_paragraph(template["description"])
            doc.add_paragraph()
            
            # 填写区域
            for field in template["fields"]:
                # 字段标签
                para = doc.add_paragraph()
                label = field["label"]
                if field.get("required"):
                    label += " *"
                para.add_run(label).bold = True
                
                if field.get("placeholder"):
                    para.add_run(f'  ({field["placeholder"]})')
                
                # 填写框
                if field["type"] == "textarea":
                    # 多行文本框
                    table = doc.add_table(rows=1, cols=1)
                    table.style = 'Table Grid'
                    cell = table.rows[0].cells[0]
                    # 添加空行作为填写空间
                    for _ in range(4):
                        cell.add_paragraph()
                elif field["type"] == "select" and field.get("options"):
                    # 选择项
                    for opt in field["options"]:
                        doc.add_paragraph(f'☐ {opt}')
                else:
                    # 单行输入
                    doc.add_paragraph('_' * 60)
                
                doc.add_paragraph()
            
            # 保存
            output_dir = output_dir or self.upload_folder
            Path(output_dir).mkdir(parents=True, exist_ok=True)
            
            filename = f'{template["title"]}.docx'
            file_path = os.path.join(output_dir, filename)
            
            doc.save(file_path)
            logger.info(f"采集表模板生成成功: {file_path}")
            
            return {
                "success": True,
                "data": {
                    "file_path": file_path,
                    "filename": filename,
                    "format": "docx"
                }
            }
            
        except Exception as e:
            logger.error(f"生成采集表模板失败: {e}")
            return {"success": False, "error": str(e)}
    
    def generate_all_templates(self, output_dir: str = None) -> Dict[str, Any]:
        """生成所有采集表模板"""
        output_dir = output_dir or os.path.join(self.upload_folder, "templates")
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        
        results = []
        for form_type in FORM_TEMPLATES:
            result = self.generate_form_template(form_type, output_dir)
            results.append({
                "form_type": form_type,
                "success": result.get("success"),
                "filename": result.get("data", {}).get("filename") if result.get("success") else None,
                "error": result.get("error")
            })
        
        success_count = sum(1 for r in results if r["success"])
        
        return {
            "success": True,
            "data": {
                "output_dir": output_dir,
                "generated_count": success_count,
                "total_count": len(FORM_TEMPLATES),
                "results": results
            }
        }
    
    def _decode_zip_filename(self, raw_filename: bytes, info) -> str:
        """
        解码zip文件名，处理中文编码问题
        支持UTF-8、GBK、GB2312编码
        """
        # 如果已经是字符串，直接返回
        if isinstance(raw_filename, str):
            # 尝试检测是否是错误解码的结果
            try:
                # 如果看起来像是cp437错误解码的中文，尝试重新解码
                if any(ord(c) > 127 for c in raw_filename):
                    try:
                        # 尝试用cp437编码回bytes再用GBK解码
                        raw_bytes = raw_filename.encode('cp437')
                        return raw_bytes.decode('utf-8')
                    except:
                        try:
                            raw_bytes = raw_filename.encode('cp437')
                            return raw_bytes.decode('gbk')
                        except:
                            pass
                return raw_filename
            except:
                return raw_filename
        
        # bytes类型，尝试多种编码
        for encoding in ['utf-8', 'gbk', 'gb2312', 'cp437']:
            try:
                return raw_filename.decode(encoding)
            except:
                continue
        
        return raw_filename.decode('utf-8', errors='replace')
    
    def process_zip_upload(self, project_id: str, zip_path: str) -> Dict[str, Any]:
        """
        处理zip文件上传：解压、分析、归类、提取内容
        """
        import zipfile
        import tempfile
        import shutil
        
        try:
            if not zipfile.is_zipfile(zip_path):
                return {"success": False, "error": "不是有效的zip文件"}
            
            # 创建临时目录解压
            temp_dir = tempfile.mkdtemp(prefix="gtv_zip_")
            extracted_files = []
            
            try:
                with zipfile.ZipFile(zip_path, 'r') as zf:
                    # 手动解压每个文件，处理中文编码
                    for info in zf.infolist():
                        # 跳过目录
                        if info.is_dir():
                            continue
                        
                        # 解码文件名
                        original_filename = info.filename
                        try:
                            # 尝试修复中文编码
                            if info.flag_bits & 0x800:
                                # UTF-8标志位已设置
                                decoded_filename = original_filename
                            else:
                                # 尝试GBK解码
                                try:
                                    decoded_filename = original_filename.encode('cp437').decode('gbk')
                                except:
                                    try:
                                        decoded_filename = original_filename.encode('cp437').decode('utf-8')
                                    except:
                                        decoded_filename = original_filename
                        except:
                            decoded_filename = original_filename
                        
                        # 获取纯文件名（不含路径）
                        filename = os.path.basename(decoded_filename)
                        
                        # 跳过隐藏文件和系统文件
                        if filename.startswith('.') or filename.startswith('~') or filename.startswith('__'):
                            continue
                        if '__MACOSX' in decoded_filename:
                            continue
                        
                        # 解压到临时目录
                        # 创建正确编码的目标路径
                        target_path = os.path.join(temp_dir, decoded_filename)
                        target_dir = os.path.dirname(target_path)
                        Path(target_dir).mkdir(parents=True, exist_ok=True)
                        
                        # 读取并写入文件
                        with zf.open(info) as source, open(target_path, 'wb') as target:
                            target.write(source.read())
                        
                        file_size = os.path.getsize(target_path)
                        file_ext = filename.split('.')[-1].lower() if '.' in filename else ''
                        
                        # 分析文件应该归类到哪里
                        category_guess = self._guess_file_category(filename, decoded_filename)
                        
                        extracted_files.append({
                            "filename": filename,
                            "relative_path": decoded_filename,
                            "temp_path": target_path,
                            "size": file_size,
                            "extension": file_ext,
                            "category_guess": category_guess
                        })
                
                # 对每个文件进行处理
                results = []
                for file_info in extracted_files:
                    result = self._process_extracted_file(project_id, file_info)
                    results.append(result)
                
                # 统计结果
                success_count = sum(1 for r in results if r.get("status") == "success")
                auto_filled = sum(1 for r in results if r.get("auto_filled"))
                unrecognized = sum(1 for r in results if r.get("status") == "unrecognized")
                
                return {
                    "success": True,
                    "data": {
                        "total_files": len(extracted_files),
                        "success_count": success_count,
                        "auto_filled_count": auto_filled,
                        "unrecognized_count": unrecognized,
                        "files": results
                    }
                }
                
            finally:
                # 清理临时目录
                shutil.rmtree(temp_dir, ignore_errors=True)
                
        except Exception as e:
            logger.error(f"处理zip文件失败: {e}")
            return {"success": False, "error": str(e)}
    
    def _guess_file_category(self, filename: str, relative_path: str = "") -> Optional[Dict[str, str]]:
        """根据文件名和路径猜测应该归类到哪个分类"""
        lower_name = filename.lower()
        lower_path = relative_path.lower()
        combined = f"{lower_path}/{lower_name}"
        
        # 简历
        if any(kw in combined for kw in ['简历', 'cv', 'resume', '履历']):
            if any(kw in combined for kw in ['推荐人', 'recommender', '推荐人1', '推荐人一']):
                if '1' in combined or '一' in combined:
                    return {"category_id": "folder_6", "item_id": "recommender_1_resume"}
                elif '2' in combined or '二' in combined:
                    return {"category_id": "folder_6", "item_id": "recommender_2_resume"}
                elif '3' in combined or '三' in combined:
                    return {"category_id": "folder_6", "item_id": "recommender_3_resume"}
            return {"category_id": "folder_1", "item_id": "resume"}
        
        # 护照
        if any(kw in combined for kw in ['护照', 'passport']):
            return {"category_id": "folder_1", "item_id": "passport"}
        
        # 学历
        if any(kw in combined for kw in ['学历', '毕业证', '学位证', 'degree', 'diploma', 'certificate', '教育']):
            return {"category_id": "folder_1", "item_id": "education"}
        
        # 专利
        if any(kw in combined for kw in ['专利', 'patent']):
            return {"category_id": "folder_4", "item_id": "patents"}
        
        # 论文
        if any(kw in combined for kw in ['论文', 'paper', 'publication', '出版']):
            return {"category_id": "folder_4", "item_id": "publications"}
        
        # 奖项
        if any(kw in combined for kw in ['奖', 'award', '荣誉', 'honor', '表彰']):
            return {"category_id": "folder_4", "item_id": "achievement_awards"}
        
        # 项目
        if any(kw in combined for kw in ['项目', 'project']):
            if any(kw in combined for kw in ['阐述', '描述', 'description', '说明']):
                return {"category_id": "folder_5", "item_id": "project_form"}
            return {"category_id": "folder_5", "item_id": "project_docs"}
        
        # 就职信息
        if any(kw in combined for kw in ['就职', '工作', 'employment', 'work']):
            if any(kw in combined for kw in ['采集', '信息', 'form']):
                return {"category_id": "folder_3", "item_id": "prev_employment_form"}
            if any(kw in combined for kw in ['证明', 'proof', 'certificate']):
                return {"category_id": "folder_3", "item_id": "prev_employment_proof"}
        
        # 推荐人
        if any(kw in combined for kw in ['推荐人', 'recommender', 'referee']):
            if '1' in combined or '一' in combined:
                return {"category_id": "folder_6", "item_id": "recommender_1_contribution_form"}
            elif '2' in combined or '二' in combined:
                return {"category_id": "folder_6", "item_id": "recommender_2_contribution_form"}
            elif '3' in combined or '三' in combined:
                return {"category_id": "folder_6", "item_id": "recommender_3_contribution_form"}
        
        # 原创贡献
        if any(kw in combined for kw in ['原创', '贡献', 'contribution', 'original']):
            return {"category_id": "folder_4", "item_id": "contribution_form"}
        
        # 收入证明
        if any(kw in combined for kw in ['收入', '工资', 'salary', 'income']):
            return {"category_id": "folder_2", "item_id": "income_proof"}
        
        # 在职证明
        if any(kw in combined for kw in ['在职', 'employment letter']):
            return {"category_id": "folder_2", "item_id": "employment_letter"}
        
        return None
    
    def _process_extracted_file(self, project_id: str, file_info: Dict[str, Any]) -> Dict[str, Any]:
        """处理解压出的单个文件"""
        import shutil
        
        result = {
            "filename": file_info["filename"],
            "size": file_info["size"],
            "extension": file_info["extension"],
            "status": "pending",
            "auto_filled": False
        }
        
        category_guess = file_info.get("category_guess")
        
        if not category_guess:
            result["status"] = "unrecognized"
            result["message"] = "无法识别文件类型，请手动上传"
            return result
        
        category_id = category_guess["category_id"]
        item_id = category_guess["item_id"]
        
        try:
            # 确保目标目录存在
            target_dir = os.path.join(self.upload_folder, project_id, category_id, item_id)
            Path(target_dir).mkdir(parents=True, exist_ok=True)
            
            # 生成唯一文件名
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_filename = f"{timestamp}_{file_info['filename']}"
            target_path = os.path.join(target_dir, safe_filename)
            
            # 复制文件
            shutil.copy2(file_info["temp_path"], target_path)
            
            # 记录到数据库
            with self._get_connection() as conn:
                cursor = conn.cursor()
                cursor.execute("""
                    INSERT INTO material_files 
                    (project_id, category_id, item_id, file_name, file_path, file_size, file_type, description)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    project_id,
                    category_id,
                    item_id,
                    file_info["filename"],
                    target_path,
                    file_info["size"],
                    file_info["extension"],
                    f"从zip自动解压: {file_info['relative_path']}"
                ))
                
                # 更新材料状态
                cursor.execute("""
                    UPDATE material_collection 
                    SET status = 'collected', collected_at = ?, file_name = ?
                    WHERE project_id = ? AND category_id = ? AND item_id = ?
                """, (
                    datetime.now().isoformat(),
                    file_info["filename"],
                    project_id,
                    category_id,
                    item_id
                ))
                
                conn.commit()
            
            result["status"] = "success"
            result["category_id"] = category_id
            result["item_id"] = item_id
            result["category_name"] = MATERIAL_CATEGORIES.get(category_id, {}).get("name", category_id)
            
            # 尝试提取文件内容进行AI分析（仅限特定类型）
            if file_info["extension"] in ['docx', 'pdf', 'txt']:
                try:
                    content = self._extract_file_content(target_path, file_info["extension"])
                    if content and len(content) > 100:
                        result["content_preview"] = content[:500] + "..." if len(content) > 500 else content
                        result["has_content"] = True
                except Exception as e:
                    logger.warning(f"提取文件内容失败: {e}")
            
            return result
            
        except Exception as e:
            logger.error(f"处理文件失败 {file_info['filename']}: {e}")
            result["status"] = "error"
            result["message"] = str(e)
            return result
    
    def _extract_file_content(self, file_path: str, extension: str) -> Optional[str]:
        """提取文件内容"""
        try:
            if extension == 'txt':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            
            elif extension == 'docx' and DOCX_AVAILABLE:
                doc = Document(file_path)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                return '\n'.join(paragraphs)
            
            elif extension == 'pdf':
                try:
                    from pdfminer.high_level import extract_text
                    return extract_text(file_path)
                except ImportError:
                    logger.warning("pdfminer未安装，无法提取PDF内容")
                    return None
            
            return None
            
        except Exception as e:
            logger.error(f"提取文件内容失败: {e}")
            return None


# 测试代码
if __name__ == "__main__":
    manager = RawMaterialManager("test_materials.db")
    
    # 测试获取分类
    categories = manager.get_material_categories()
    print(f"材料分类数: {len(categories['data'])}")
    
    # 测试初始化项目
    result = manager.init_project_materials("TEST001")
    print(f"初始化结果: {result}")
    
    # 测试获取状态
    status = manager.get_collection_status("TEST001")
    print(f"进度: {status['data']['progress']}")
    
    # 测试完整性检查
    check = manager.check_completeness("TEST001")
    print(f"必填缺失: {check['data']['summary']['required_missing_count']}")
