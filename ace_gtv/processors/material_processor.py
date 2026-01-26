#!/usr/bin/env python3
"""
材料处理器 - 处理各种类型的原始材料
支持：PDF、Word、图片、网站链接等
"""

import os
import re
import json
import base64
import tempfile
import mimetypes
from datetime import datetime
from typing import Dict, Any, List, Optional, Tuple
from pathlib import Path
from urllib.parse import urlparse
import requests

from utils.logger_config import setup_module_logger

logger = setup_module_logger("material_processor", os.getenv("LOG_LEVEL", "INFO"))

# 可选导入
try:
    from pdfminer.high_level import extract_text as pdf_extract_text
    from pdfminer.pdfparser import PDFSyntaxError
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    logger.warning("pdfminer.six未安装，PDF处理功能不可用")

try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    logger.warning("python-docx未安装，Word文档处理功能不可用")

try:
    from PIL import Image
    import io
    PIL_SUPPORT = True
except ImportError:
    PIL_SUPPORT = False
    logger.warning("Pillow未安装，图片处理功能受限")

try:
    from bs4 import BeautifulSoup
    BS4_SUPPORT = True
except ImportError:
    BS4_SUPPORT = False
    logger.warning("beautifulsoup4未安装，网页解析功能受限")

try:
    from openai import OpenAI
    OPENAI_SUPPORT = True
except ImportError:
    OPENAI_SUPPORT = False


class MaterialProcessor:
    """材料处理器 - 处理和提取各种格式的内容"""
    
    # 支持的文件类型
    SUPPORTED_EXTENSIONS = {
        # 文档类
        'pdf': 'document',
        'doc': 'document',
        'docx': 'document',
        'txt': 'text',
        'md': 'text',
        'rtf': 'document',
        
        # 图片类
        'jpg': 'image',
        'jpeg': 'image',
        'png': 'image',
        'gif': 'image',
        'webp': 'image',
        'bmp': 'image',
        'tiff': 'image',
        'heic': 'image',
        
        # 数据类
        'json': 'data',
        'csv': 'data',
        'xlsx': 'data',
        'xls': 'data',
    }
    
    # 材料分类
    MATERIAL_CATEGORIES = {
        "简历": ["resume", "cv", "简历", "履历"],
        "推荐信": ["recommendation", "reference", "推荐", "endorsement"],
        "证书": ["certificate", "diploma", "证书", "学位", "资质"],
        "论文": ["paper", "publication", "论文", "文章", "研究"],
        "专利": ["patent", "专利", "发明"],
        "奖项": ["award", "prize", "奖", "荣誉"],
        "媒体报道": ["news", "media", "press", "报道", "新闻"],
        "作品集": ["portfolio", "project", "作品", "项目"],
        "个人陈述": ["statement", "personal", "陈述", "自述"],
        "其他": []
    }
    
    def __init__(self, upload_folder: str = "./uploads"):
        """
        初始化材料处理器
        
        Args:
            upload_folder: 上传文件存储目录
        """
        self.upload_folder = Path(upload_folder)
        self.upload_folder.mkdir(parents=True, exist_ok=True)
        
        # 初始化OpenAI客户端（用于图片OCR）
        self.openai_client = None
        if OPENAI_SUPPORT:
            api_key = os.getenv("OPENAI_API_KEY")
            if api_key:
                self.openai_client = OpenAI(api_key=api_key)
        
        logger.info(f"材料处理器初始化完成，上传目录: {self.upload_folder}")
        logger.info(f"功能支持: PDF={PDF_SUPPORT}, DOCX={DOCX_SUPPORT}, PIL={PIL_SUPPORT}, BS4={BS4_SUPPORT}")
    
    def process_material(self, file_data: bytes, filename: str, 
                        category: str = None) -> Dict[str, Any]:
        """
        处理上传的材料文件
        
        Args:
            file_data: 文件二进制数据
            filename: 文件名
            category: 材料分类（可选，会自动推断）
            
        Returns:
            处理结果字典
        """
        try:
            # 获取文件扩展名和类型
            ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
            file_type = self.SUPPORTED_EXTENSIONS.get(ext, 'unknown')
            
            # 保存原始文件
            saved_path = self._save_file(file_data, filename)
            
            # 提取内容
            content = ""
            metadata = {
                "filename": filename,
                "extension": ext,
                "file_type": file_type,
                "file_size": len(file_data),
                "saved_path": str(saved_path),
                "processed_at": datetime.now().isoformat()
            }
            
            if file_type == 'document':
                if ext == 'pdf':
                    content, pdf_meta = self._extract_pdf(saved_path)
                    metadata.update(pdf_meta)
                elif ext in ['doc', 'docx']:
                    content, doc_meta = self._extract_docx(saved_path)
                    metadata.update(doc_meta)
                elif ext == 'rtf':
                    content = self._extract_rtf(saved_path)
                    
            elif file_type == 'text':
                content = file_data.decode('utf-8', errors='ignore')
                
            elif file_type == 'image':
                content, img_meta = self._extract_image(saved_path, file_data)
                metadata.update(img_meta)
                
            elif file_type == 'data':
                content, data_meta = self._extract_data_file(saved_path, ext)
                metadata.update(data_meta)
            
            # 自动分类
            if not category:
                category = self._auto_categorize(filename, content)
            
            metadata["category"] = category
            metadata["content_length"] = len(content)
            
            # 提取关键信息摘要
            summary = self._generate_summary(content, file_type)
            
            return {
                "success": True,
                "filename": filename,
                "category": category,
                "file_type": file_type,
                "content": content,
                "summary": summary,
                "metadata": metadata,
                "saved_path": str(saved_path)
            }
            
        except Exception as e:
            logger.error(f"处理材料失败 {filename}: {e}")
            return {
                "success": False,
                "error": str(e),
                "filename": filename
            }
    
    def process_url(self, url: str, category: str = None) -> Dict[str, Any]:
        """
        处理网站链接，提取内容
        
        Args:
            url: 网站URL
            category: 材料分类
            
        Returns:
            处理结果字典
        """
        try:
            # 验证URL
            parsed = urlparse(url)
            if not parsed.scheme or not parsed.netloc:
                return {"success": False, "error": "无效的URL格式"}
            
            # 获取网页内容
            headers = {
                'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36'
            }
            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()
            
            # 解析内容
            content_type = response.headers.get('content-type', '')
            
            if 'text/html' in content_type:
                content, metadata = self._parse_html(response.text, url)
            elif 'application/json' in content_type:
                content = json.dumps(response.json(), ensure_ascii=False, indent=2)
                metadata = {"type": "json"}
            elif 'application/pdf' in content_type:
                # 下载PDF并处理
                return self.process_material(response.content, f"download_{parsed.netloc}.pdf", category)
            else:
                content = response.text
                metadata = {"type": "text"}
            
            # 自动分类
            if not category:
                category = self._auto_categorize(url, content)
            
            metadata.update({
                "url": url,
                "domain": parsed.netloc,
                "processed_at": datetime.now().isoformat(),
                "category": category
            })
            
            # 生成摘要
            summary = self._generate_summary(content, "webpage")
            
            return {
                "success": True,
                "url": url,
                "category": category,
                "file_type": "webpage",
                "content": content,
                "summary": summary,
                "metadata": metadata
            }
            
        except requests.RequestException as e:
            logger.error(f"获取URL内容失败 {url}: {e}")
            return {"success": False, "error": f"无法访问URL: {str(e)}", "url": url}
        except Exception as e:
            logger.error(f"处理URL失败 {url}: {e}")
            return {"success": False, "error": str(e), "url": url}
    
    def _save_file(self, file_data: bytes, filename: str) -> Path:
        """保存文件到上传目录"""
        # 创建日期子目录
        date_folder = self.upload_folder / datetime.now().strftime("%Y%m%d")
        date_folder.mkdir(parents=True, exist_ok=True)
        
        # 处理文件名冲突
        base_name = Path(filename).stem
        ext = Path(filename).suffix
        save_path = date_folder / filename
        
        counter = 1
        while save_path.exists():
            save_path = date_folder / f"{base_name}_{counter}{ext}"
            counter += 1
        
        save_path.write_bytes(file_data)
        logger.info(f"文件保存成功: {save_path}")
        return save_path
    
    def _extract_pdf(self, file_path: Path) -> Tuple[str, Dict]:
        """提取PDF内容"""
        metadata = {"pdf_pages": 0}
        
        if not PDF_SUPPORT:
            return "[PDF处理功能未启用]", metadata
        
        try:
            content = pdf_extract_text(str(file_path))
            
            # 获取页数
            try:
                from pdfminer.pdfpage import PDFPage
                with open(file_path, 'rb') as f:
                    pages = list(PDFPage.get_pages(f))
                    metadata["pdf_pages"] = len(pages)
            except:
                pass
            
            return content.strip(), metadata
            
        except PDFSyntaxError as e:
            logger.warning(f"PDF语法错误: {e}")
            return "[PDF解析失败，可能是加密或损坏的文件]", metadata
        except Exception as e:
            logger.error(f"PDF提取失败: {e}")
            return f"[PDF提取错误: {str(e)}]", metadata
    
    def _extract_docx(self, file_path: Path) -> Tuple[str, Dict]:
        """提取Word文档内容"""
        metadata = {"doc_paragraphs": 0, "doc_tables": 0}
        
        if not DOCX_SUPPORT:
            return "[Word文档处理功能未启用]", metadata
        
        try:
            doc = DocxDocument(str(file_path))
            
            # 提取段落
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            # 提取表格
            tables_content = []
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                tables_content.append("\n".join(table_text))
            
            metadata["doc_paragraphs"] = len(paragraphs)
            metadata["doc_tables"] = len(doc.tables)
            
            content = "\n\n".join(paragraphs)
            if tables_content:
                content += "\n\n--- 表格内容 ---\n\n" + "\n\n".join(tables_content)
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"Word文档提取失败: {e}")
            return f"[Word文档提取错误: {str(e)}]", metadata
    
    def _extract_rtf(self, file_path: Path) -> str:
        """提取RTF文档内容"""
        try:
            content = file_path.read_text(encoding='utf-8', errors='ignore')
            # 简单的RTF标签清理
            content = re.sub(r'\\[a-z]+\d*\s?', '', content)
            content = re.sub(r'[{}]', '', content)
            return content.strip()
        except Exception as e:
            return f"[RTF提取错误: {str(e)}]"
    
    def _extract_image(self, file_path: Path, file_data: bytes) -> Tuple[str, Dict]:
        """提取图片内容（OCR）"""
        metadata = {"image_width": 0, "image_height": 0, "ocr_method": None}
        
        # 获取图片尺寸
        if PIL_SUPPORT:
            try:
                img = Image.open(io.BytesIO(file_data))
                metadata["image_width"] = img.width
                metadata["image_height"] = img.height
                metadata["image_mode"] = img.mode
            except Exception as e:
                logger.warning(f"读取图片信息失败: {e}")
        
        # 使用OpenAI Vision API进行OCR
        if self.openai_client:
            try:
                base64_image = base64.b64encode(file_data).decode('utf-8')
                
                # 确定MIME类型
                ext = file_path.suffix.lower().lstrip('.')
                mime_type = mimetypes.guess_type(str(file_path))[0] or f"image/{ext}"
                
                response = self.openai_client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "text",
                                    "text": """请仔细查看这张图片，完成以下任务：
1. 如果图片包含文字，请完整提取所有文字内容
2. 如果图片是证书、奖状、推荐信等文档，请提取所有关键信息
3. 如果图片是简历或履历，请提取个人信息、教育背景、工作经历等
4. 描述图片中的任何重要视觉元素

请用原文语言输出提取的内容，保持原有格式。"""
                                },
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:{mime_type};base64,{base64_image}"
                                    }
                                }
                            ]
                        }
                    ],
                    max_tokens=4096
                )
                
                content = response.choices[0].message.content
                metadata["ocr_method"] = "openai_vision"
                return content, metadata
                
            except Exception as e:
                logger.error(f"OpenAI Vision OCR失败: {e}")
                metadata["ocr_error"] = str(e)
        
        # 如果OpenAI不可用，返回占位内容
        return f"[图片文件: {file_path.name}，需要手动查看或配置OCR服务]", metadata
    
    def _extract_data_file(self, file_path: Path, ext: str) -> Tuple[str, Dict]:
        """提取数据文件内容"""
        metadata = {"data_type": ext}
        
        try:
            if ext == 'json':
                data = json.loads(file_path.read_text(encoding='utf-8'))
                content = json.dumps(data, ensure_ascii=False, indent=2)
                metadata["json_keys"] = list(data.keys()) if isinstance(data, dict) else None
                
            elif ext == 'csv':
                content = file_path.read_text(encoding='utf-8', errors='ignore')
                lines = content.split('\n')
                metadata["csv_rows"] = len(lines)
                metadata["csv_headers"] = lines[0] if lines else None
                
            elif ext in ['xlsx', 'xls']:
                try:
                    import pandas as pd
                    df = pd.read_excel(file_path)
                    content = df.to_string()
                    metadata["excel_rows"] = len(df)
                    metadata["excel_columns"] = list(df.columns)
                except ImportError:
                    content = "[Excel处理需要pandas库]"
                    
            else:
                content = file_path.read_text(encoding='utf-8', errors='ignore')
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"数据文件提取失败: {e}")
            return f"[数据文件提取错误: {str(e)}]", metadata
    
    def _parse_html(self, html: str, url: str) -> Tuple[str, Dict]:
        """解析HTML内容"""
        metadata = {"html_title": None, "html_description": None}
        
        if not BS4_SUPPORT:
            # 简单的HTML标签清理
            content = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL)
            content = re.sub(r'<style[^>]*>.*?</style>', '', content, flags=re.DOTALL)
            content = re.sub(r'<[^>]+>', ' ', content)
            content = re.sub(r'\s+', ' ', content).strip()
            return content, metadata
        
        try:
            soup = BeautifulSoup(html, 'html.parser')
            
            # 提取元数据
            if soup.title:
                metadata["html_title"] = soup.title.string
            
            meta_desc = soup.find('meta', attrs={'name': 'description'})
            if meta_desc:
                metadata["html_description"] = meta_desc.get('content')
            
            # 移除不需要的元素
            for element in soup(['script', 'style', 'nav', 'footer', 'header', 'aside']):
                element.decompose()
            
            # 提取主要内容
            main_content = soup.find('main') or soup.find('article') or soup.find('body')
            
            if main_content:
                # 提取文本
                text_parts = []
                for element in main_content.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'li', 'td', 'th']):
                    text = element.get_text(strip=True)
                    if text and len(text) > 10:  # 过滤太短的文本
                        if element.name.startswith('h'):
                            text_parts.append(f"\n## {text}\n")
                        else:
                            text_parts.append(text)
                
                content = "\n".join(text_parts)
            else:
                content = soup.get_text(separator='\n', strip=True)
            
            # 清理多余空白
            content = re.sub(r'\n{3,}', '\n\n', content)
            
            return content.strip(), metadata
            
        except Exception as e:
            logger.error(f"HTML解析失败: {e}")
            return html, metadata
    
    def _auto_categorize(self, filename: str, content: str) -> str:
        """自动分类材料"""
        text = (filename + " " + content[:2000]).lower()
        
        for category, keywords in self.MATERIAL_CATEGORIES.items():
            if category == "其他":
                continue
            for keyword in keywords:
                if keyword.lower() in text:
                    return category
        
        return "其他"
    
    def _generate_summary(self, content: str, file_type: str, max_length: int = 500) -> str:
        """生成内容摘要"""
        if not content or len(content) < 100:
            return content
        
        # 简单的摘要：取前N个字符
        summary = content[:max_length]
        if len(content) > max_length:
            # 尝试在句子边界截断
            last_period = summary.rfind('。')
            last_dot = summary.rfind('.')
            cut_point = max(last_period, last_dot)
            if cut_point > max_length // 2:
                summary = summary[:cut_point + 1]
            else:
                summary += "..."
        
        return summary
    
    def batch_process(self, materials: List[Dict]) -> Dict[str, Any]:
        """
        批量处理材料
        
        Args:
            materials: 材料列表，每个材料是字典 {type: 'file'|'url', data: ..., filename: ..., category: ...}
            
        Returns:
            批量处理结果
        """
        results = []
        success_count = 0
        
        for material in materials:
            material_type = material.get('type', 'file')
            
            if material_type == 'url':
                result = self.process_url(
                    material.get('url'),
                    material.get('category')
                )
            else:
                result = self.process_material(
                    material.get('data'),
                    material.get('filename', 'unknown'),
                    material.get('category')
                )
            
            results.append(result)
            if result.get('success'):
                success_count += 1
        
        return {
            "success": True,
            "total": len(materials),
            "success_count": success_count,
            "failed_count": len(materials) - success_count,
            "results": results
        }
    
    def get_supported_types(self) -> Dict[str, List[str]]:
        """获取支持的文件类型"""
        types = {"document": [], "text": [], "image": [], "data": []}
        for ext, file_type in self.SUPPORTED_EXTENSIONS.items():
            types[file_type].append(ext)
        return types
    
    def consolidate_materials(self, materials: List[Dict]) -> Dict[str, Any]:
        """
        整合所有材料，生成综合报告
        
        Args:
            materials: 处理后的材料列表
            
        Returns:
            整合报告
        """
        consolidated = {
            "total_materials": len(materials),
            "by_category": {},
            "by_type": {},
            "all_content": [],
            "summary_report": ""
        }
        
        for material in materials:
            if not material.get('success'):
                continue
            
            category = material.get('category', '其他')
            file_type = material.get('file_type', 'unknown')
            
            # 按分类统计
            if category not in consolidated["by_category"]:
                consolidated["by_category"][category] = []
            consolidated["by_category"][category].append({
                "filename": material.get('filename') or material.get('url'),
                "summary": material.get('summary'),
                "content_length": len(material.get('content', ''))
            })
            
            # 按类型统计
            if file_type not in consolidated["by_type"]:
                consolidated["by_type"][file_type] = 0
            consolidated["by_type"][file_type] += 1
            
            # 收集所有内容
            consolidated["all_content"].append({
                "source": material.get('filename') or material.get('url'),
                "category": category,
                "content": material.get('content')
            })
        
        # 生成摘要报告
        report_lines = [
            "# 材料整理报告",
            f"\n处理日期: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
            f"\n总计材料: {len(materials)} 份",
            "\n## 材料分类统计\n"
        ]
        
        for category, items in consolidated["by_category"].items():
            report_lines.append(f"- **{category}**: {len(items)} 份")
            for item in items:
                report_lines.append(f"  - {item['filename']}")
        
        report_lines.append("\n## 文件类型统计\n")
        for file_type, count in consolidated["by_type"].items():
            report_lines.append(f"- {file_type}: {count} 份")
        
        consolidated["summary_report"] = "\n".join(report_lines)
        
        return consolidated


# 测试代码
if __name__ == "__main__":
    processor = MaterialProcessor()
    
    # 测试URL处理
    result = processor.process_url("https://example.com")
    print(f"URL处理结果: {result.get('success')}")
    
    # 显示支持的类型
    types = processor.get_supported_types()
    print(f"支持的文件类型: {types}")
